"""Create COMP4431_Full_Exam_Notes.docx in the same format as DSAI4205_Exam_Review.docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.8); section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ── colour palette (mirrors 4205 doc) ──
C_BLUE   = RGBColor(0x25,0x63,0xEB)   # main content
C_GREEN  = RGBColor(0x10,0xB9,0x81)   # ELI5
C_AMBER  = RGBColor(0xF5,0x9E,0x0B)   # NOTE label
C_AMBERT = RGBColor(0x92,0x40,0x0E)   # NOTE text
C_CYAN   = RGBColor(0x06,0xB6,0xD4)   # formulas / pseudocode
C_PURP   = RGBColor(0xA7,0x8B,0xFA)   # code blocks
C_DPURP  = RGBColor(0x7C,0x3A,0xED)   # highlighted bullets
C_H1     = RGBColor(0x1F,0x49,0x7D)
C_H2     = RGBColor(0x2E,0x74,0xB5)

def h1(t):
    p = doc.add_heading(t, 1); p.runs[0].font.color.rgb = C_H1; return p
def h3(t):
    p = doc.add_heading(t, 3); return p
def h4(t):
    p = doc.add_heading(t, 4); return p
def eli5(t):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('ELI5: '); r.bold = True; r.font.color.rgb = C_GREEN; r.font.size = Pt(10)
    r2 = p.add_run(t); r2.font.color.rgb = C_GREEN; r2.font.size = Pt(10); return p
def blue(t, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix); r1.bold = True; r1.font.color.rgb = C_BLUE; r1.font.size = Pt(10)
    r2 = p.add_run(t); r2.font.color.rgb = C_BLUE; r2.font.size = Pt(10); return p
def note(t):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('NOTE: '); r1.bold = True; r1.font.color.rgb = C_AMBER; r1.font.size = Pt(10)
    r2 = p.add_run(t); r2.font.color.rgb = C_AMBERT; r2.font.size = Pt(10); return p
def formula(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t); r.font.color.rgb = C_CYAN; r.font.name = 'Courier New'; r.font.size = Pt(10); return p
def bullet(t, highlight=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(t); r.font.size = Pt(10)
    if highlight: r.font.color.rgb = C_DPURP
    return p
def sp(): doc.add_paragraph().paragraph_format.space_after = Pt(0)
def tbl(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'
    hrow = t.rows[0]
    for i,h in enumerate(headers):
        c = hrow.cells[i]; c.text = h
        c.paragraphs[0].runs[0].bold = True; c.paragraphs[0].runs[0].font.size = Pt(9)
        sh = OxmlElement('w:shd'); sh.set(qn('w:fill'),'D6E4F0'); sh.set(qn('w:val'),'clear')
        c._tc.get_or_add_tcPr().append(sh)
    for ri,row in enumerate(rows):
        tr = t.rows[ri+1]
        for ci,val in enumerate(row):
            c = tr.cells[ci]; c.text = val; c.paragraphs[0].runs[0].font.size = Pt(9)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    sp()

# ══════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════
title = doc.add_heading('COMP4431 Artificial Intelligence — Full Exam Notes (with ELI5)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('All content from L01–L09 | Past Paper Solutions | Quick Reference Summary')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Green italic text = ELI5 (Explain Like I\'m 5) — plain-English analogy for every concept')
r.font.color.rgb = C_GREEN; r.font.size = Pt(10)
doc.add_page_break()

# ══════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════
h1('Table of Contents')
toc_items = [
    'L01: Introduction to Artificial Intelligence',
    '  Agent Types and Environments','  Discriminative vs Generative AI','  7 Dartmouth Research Directions',
    'L02: Heuristic Search',
    '  BFS — Breadth-First Search','  DFS — Depth-First Search','  UCS — Uniform Cost Search',
    '  Greedy Search','  A* Search','  Admissibility & Heuristic Dominance',
    'L03: Knowledge-Based Agents',
    '  Forward Chaining (FC)','  Backward Chaining (BC)','  FC vs BC Comparison',
    'L04: Local Search',
    '  Hill Climbing','  Simulated Annealing (SA)','  Genetic Algorithm (GA)',
    '  GA for Feature Selection',
    'L05: Constraint Satisfaction & Games',
    '  CSP — Constraint Satisfaction Problem','  MCV, LCV, Degree Heuristics',
    '  Minimax','  Alpha-Beta Pruning',
    'L06: Machine Learning',
    '  Supervised vs Unsupervised','  K-Means Clustering',
    '  Decision Tree & Entropy','  SVM — Support Vector Machine',
    'L07: Neural Networks',
    '  Biological vs Artificial Neuron','  Perceptron Learning',
    '  MLP — Multi-Layer Perceptron','  Backpropagation','  Gradient Descent Variants',
    'L08: Deep Learning',
    '  CNN — Convolutional Neural Network','  CNN Convolution Calculation',
    '  RNN — Recurrent Neural Network','  LSTM','  GAN',
    'L09: Bayesian Networks',
    '  Probability & Bayes Rule','  Bayesian Network Inference',
    '  Naive Bayes Classifier',
    'Quick Reference Summary',
    '  Must-Know Formulas','  Key Comparisons at a Glance','  30-Second Calculation Recipes',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(item); r.font.color.rgb = C_BLUE; r.font.size = Pt(10)
doc.add_page_break()

# ══════════════════════════════════════════════
h1('L01: Introduction to Artificial Intelligence')
# ══════════════════════════════════════════════

h3('What is AI?')
eli5('Teaching a computer to do smart things — learn, reason, plan — instead of programming every single step by hand.')
sp()
bullet('AI = systems that perceive environment and take actions to maximise a performance measure.')
bullet('Rational agent: acts to achieve the best expected outcome given its percepts.')
bullet('Strong AI: human-level general intelligence (still theoretical).')
bullet('Weak AI: narrow task-specific intelligence — what we build today.')
sp()

h3('Agent Types and Environments')
eli5('A thermostat is a dumb agent — it only senses temperature and flips a switch. A chess program is smarter — it looks ahead and plans moves.')
sp()
tbl(['Agent Type','Memory','Modelling','Example'],
    [['Simple Reflex','None','None','Thermostat'],
     ['Model-Based Reflex','Internal state','Environment model','Roomba vacuum'],
     ['Goal-Based','Internal state','Goals','Route-planning GPS'],
     ['Utility-Based','Internal state','Goals + preferences','Self-driving car'],
     ['Learning Agent','All of above + learning','Adapts over time','AlphaGo, ChatGPT'],
    ],widths=[1.6,1.2,1.4,2.2])
note('Exam tip: Identify agent type from a short description. Key question: does it use memory? Does it model the world? Does it have preferences?')

h3('Discriminative vs Generative AI')
eli5('Discriminative = learns the fence between cats and dogs. Generative = learns what a cat looks like so it can draw one from scratch.')
sp()
tbl(['','Discriminative','Generative'],
    [['Learns','Decision boundary P(Y|X)','Data distribution P(X,Y) or P(X)'],
     ['Goal','Classify or label input','Generate new samples / understand data'],
     ['Examples','SVM, Logistic Reg, CNN classifier','GAN, VAE, Naive Bayes, RBM'],
    ],widths=[1.4,2.6,2.6])
note('Exam tip: GAN and VAE = generative. SVM and CNN (classification mode) = discriminative.')

h3('7 Dartmouth Research Directions (1956)')
eli5('At a famous 1956 summer camp, scientists wrote down the 7 things they wanted AI to do. These became the map for all of AI research.')
sp()
tbl(['#','Direction','Plain English'],
    [['1','Automatic Computers','Machines that compute on their own'],
     ['2','Programming Language for AI','Teach machines using language, not just numbers'],
     ['3','Neuron Nets','Build brains from connected nodes — leads to deep learning'],
     ['4','Theory of the Size of a Calculation','How hard is a problem really?'],
     ['5','Self-Improvement (Learning)','Machines that get better by themselves'],
     ['6','Abstractions','Handle concepts, not just raw data'],
     ['7','Randomness & Creativity','Use chance to generate novel solutions'],
    ],widths=[0.4,2.4,3.6])
sp()

doc.add_page_break()

# ══════════════════════════════════════════════
h1('L02: Heuristic Search')
# ══════════════════════════════════════════════

h3('Search Problem Setup')
eli5('You are in a maze. You need to find the exit. The search algorithm decides WHICH door to open next.')
sp()
tbl(['Component','Meaning','Example (Romania map)'],
    [['State','A snapshot of the world','Current city'],
     ['Initial state','Where you start','Arad'],
     ['Goal test','Are we done?','city == Bucharest'],
     ['Actions','Possible moves','Drive to adjacent city'],
     ['Path cost','Cost of a sequence of actions','Total km driven'],
     ['Heuristic h(n)','Estimated cost to goal from n','Straight-line distance to Bucharest'],
    ],widths=[1.4,2.4,2.6])

h3('BFS — Breadth-First Search')
eli5('Try every door on floor 1 before going to floor 2. Guaranteed fewest hops, but ignores corridor length.')
sp()
blue('Step 1: ','Add start node to FIFO queue. Mark visited.')
blue('Step 2: ','Dequeue node n. If n is goal → done.')
blue('Step 3: ','Add all unvisited neighbours to queue (alphabetical if tie-breaking specified). Mark visited.')
blue('Step 4: ','Repeat until goal found or queue empty.')
sp()
tbl(['Property','BFS'],
    [['Complete?','Yes (finite graph)'],
     ['Optimal?','Yes — fewest hops (not cheapest cost)'],
     ['Time','O(b^d)  b=branching factor, d=depth of solution'],
     ['Space','O(b^d) — stores entire frontier'],
    ],widths=[1.8,4.6])
note('Exam tip: BFS finds fewest edges, NOT cheapest path. A→B→D→E (3 hops) loses to A→C→E (2 hops) even if A→C→E costs more.')

h3('DFS — Depth-First Search')
eli5('Pick one corridor, run all the way to the dead end, backtrack, try next. Fast but might take a terrible detour.')
sp()
blue('Rule: ','Use a LIFO stack (or recursion). Always expand the deepest node first.')
tbl(['Property','DFS'],
    [['Complete?','No — may loop in infinite graphs'],
     ['Optimal?','No — finds first path, not best'],
     ['Time','O(b^m)  m=max depth'],
     ['Space','O(bm) — only stores current path'],
    ],widths=[1.8,4.6])
note('Exam tip: DFS uses FAR less memory than BFS. In exams, trace alphabetically — A expands B before C if both neighbours.')

h3('UCS — Uniform Cost Search')
eli5('Like BFS but you care about corridor length. Always expand the cheapest path so far. Finds the cheapest total route.')
sp()
blue('Rule: ','Priority queue ordered by g(n) = cumulative path cost from start. Expand lowest g(n) first.')
tbl(['Property','UCS'],
    [['Complete?','Yes (positive step costs)'],
     ['Optimal?','Yes — finds minimum cost path'],
     ['Key difference from BFS','BFS counts hops; UCS counts cost'],
    ],widths=[1.8,4.6])
note('Exam tip: UCS = Dijkstra\'s algorithm. Update a node if you find a cheaper path to it.')

h3('Greedy Search')
eli5('Always walk toward wherever smells most like the exit. Fast but can walk into dead ends.')
sp()
blue('Rule: ','Priority queue ordered by h(n) = heuristic estimate to goal. Ignores g(n) entirely.')
note('Exam tip: Greedy is NOT optimal and NOT complete. It can get stuck if h is misleading.')

h3('A* Search')
eli5('Best of both: actual distance walked so far + smart estimate of distance left. Never overestimates. Finds the shortest route guaranteed.')
sp()
blue('Rule: ','Priority queue ordered by f(n) = g(n) + h(n). Expand lowest f first.')
blue('Update: ','If a cheaper path to an already-seen node is found, update its g(n) and re-queue.')
sp()
formula('f(n) = g(n) + h(n)')
formula('g(n) = cost from start to n (exact)')
formula('h(n) = estimated cost from n to goal (heuristic)')
sp()
tbl(['Step','Action'],
    [['1','Initialise open list with start node. f(start)=0+h(start).'],
     ['2','Pop node with lowest f from open list.'],
     ['3','If it is the goal → reconstruct path and stop.'],
     ['4','For each neighbour: compute g_new = g(current)+edge_cost. If cheaper than known g, update and add to open list.'],
     ['5','Repeat until goal popped or open list empty.'],
    ],widths=[0.4,6.0])
note('Exam tip: Always show the open list with f values at each step. When expanding, list ALL frontier nodes and their f values.')

h4('Dec 2024 Q4 — A* Timisoara to Bucharest (Romania map) [12 marks — full expansion required]')
blue('SLD heuristics: ','Timisoara=329, Lugoj=244, Mehadia=241, Arad=366, Dobreta=242, Sibiu=253, RimnicuVilcea=193, Fagaras=176, Pitesti=100, Bucharest=0')
blue('Edge costs: ','Tim-Arad=118, Tim-Lugoj=111, Lugoj-Mehadia=70, Mehadia-Dobreta=75, Dobreta-Craiova=120, Arad-Sibiu=140, Sibiu-RV=80, Sibiu-Fagaras=99, RV-Pitesti=97, Pitesti-Bucharest=101')
tbl(['Step','Expanded','g','h','f','Frontier (key nodes)'],
    [['1','Timisoara','0','329','329','Lugoj(355), Arad(484)'],
     ['2','Lugoj','111','244','355','Mehadia(422), Arad(484)'],
     ['3','Mehadia','181','241','422','Arad(484), Dobreta(498)'],
     ['4','Arad','118','366','484','Dobreta(498), Sibiu(511)'],
     ['5','Dobreta','256','242','498','Sibiu(511), Craiova(536)'],
     ['6','Sibiu','258','253','511','RimnicuVilcea(531), Fagaras(533), Craiova(536)'],
     ['7','RimnicuVilcea','338','193','531','Fagaras(533), Pitesti(535), Craiova(536)'],
     ['8','Fagaras','357','176','533','Pitesti(535), Craiova(536), Bucharest_Fagaras(568)'],
     ['9','Pitesti','435','100','535','Craiova(536), Bucharest_Pitesti(536)←replaces 568'],
     ['10','Bucharest','536','0','536','GOAL'],
    ],widths=[0.4,1.3,0.6,0.6,0.7,3.2])
formula('Path: Timisoara → Arad → Sibiu → RimnicuVilcea → Pitesti → Bucharest   Cost = 118+140+80+97+101 = 536 km')
note('Arad(g=118,h=366,f=484) is expanded BEFORE Dobreta(f=498) even though Arad is not on the optimal path. A* must explore it because its f is lower at that point. Fagaras route to Bucharest (g=568) is discarded when Pitesti finds cheaper g=536.')

h3('Admissibility & Heuristic Dominance')
eli5('Admissible = your estimate of how far to the exit is never too HIGH. You can be too LOW (optimistic), never too high. This guarantees A* finds the best path.')
sp()
blue('Admissible: ','h(n) <= h*(n) for all n, where h*(n) is the true cost to goal. Never overestimates.')
blue('Consistent (monotone): ','h(n) <= cost(n,n\')+h(n\') for all neighbours n\'. Ensures A* never re-expands nodes.')
blue('Dominant: ','h2 dominates h1 if h2(n)>=h1(n) for all n AND both are admissible. Dominant = more informed = fewer nodes expanded.')
sp()
tbl(['Heuristic','Description','Admissible?','Dominates'],
    [['h1 = misplaced tiles','Count tiles not in goal position','Yes','No'],
     ['h2 = Manhattan distance','Sum of |row-diff|+|col-diff| for each tile','Yes','Yes — h2>=h1 always'],
     ['h = 0','Trivial — expands everything','Yes','No (weakest possible)'],
    ],widths=[1.8,2.4,1.1,1.9])
note('Exam tip: To prove admissibility for a node, show h(n) <= h*(n) using the actual shortest path from that node to the goal.')

doc.add_page_break()

# ══════════════════════════════════════════════
h1('L03: Knowledge-Based Agents')
# ══════════════════════════════════════════════

h3('Knowledge Base (KB) and Inference')
eli5('A detective has a notebook of facts and rules. Inference = using those rules to figure out new facts the detective didn\'t directly observe.')
sp()
bullet('KB = set of sentences in a formal language (propositional or first-order logic).')
bullet('TELL: add a new fact to the KB.')
bullet('ASK: query the KB to check if something is true.')
bullet('Inference: derive new sentences that are entailed by existing sentences.')
sp()
note('Exam tip: "Query adds sentences to KB" → FALSE. TELL adds; ASK queries.')

h3('Forward Chaining (FC)')
eli5('Start from clues you HAVE. Apply rules. Write down new conclusions. Repeat until the goal is proved or no more rules fire. Bottom-up.')
sp()
blue('Step 1: ','Begin with all known facts in KB.')
blue('Step 2: ','Find any rule whose ALL premises are already satisfied.')
blue('Step 3: ','Fire that rule — add its conclusion to KB as a new fact.')
blue('Step 4: ','Repeat until goal is in KB (success) or no new facts can be derived (failure).')
sp()
h4('FC Example — Car Diagnosis')
tbl(['Step','Rule fired','New fact added'],
    [['1','headlights_dim → battery_dead','battery_dead'],
     ['2','car_wont_start AND battery_dead → dead_battery_issue','dead_battery_issue'],
    ],widths=[0.5,3.0,2.9])
note('Exam tip: FC = data-driven / bottom-up. Good when many queries needed or all consequences required.')

h3('Backward Chaining (BC)')
eli5('Start from what you WANT to prove. Work backwards to see if the premises are satisfied. Top-down.')
sp()
blue('Step 1: ','Start with the goal query.')
blue('Step 2: ','Find rules that CONCLUDE the goal.')
blue('Step 3: ','Their premises become new sub-goals. Recursively prove each sub-goal.')
blue('Step 4: ','Succeed if all sub-goals are in KB. Fail if any sub-goal is unprovable.')
sp()
h4('BC Example — KB: {A,B,C, A∧B→D, B∧C→E, A∧C→F, A∧F→G, D∧F→K, G∧K→Q1, H∧C→Q2}')
tbl(['Goal','BC trace','Result'],
    [['Q1','Q1←G∧K; G←A∧F; F←A∧C ✓; K←D∧F; D←A∧B ✓','PROVED'],
     ['Q2','Q2←H∧C; C given, H not in KB and not derivable','FAILS'],
    ],widths=[0.8,4.4,0.9])
note('Exam tip: BC = goal-driven / top-down. Good for single specific goal — only explores what is needed.')

h3('FC vs BC Comparison')
tbl(['','Forward Chaining','Backward Chaining'],
    [['Direction','Facts → Conclusions (bottom-up)','Goal → Premises (top-down)'],
     ['Start point','Known facts','Goal query'],
     ['Derives','ALL consequences','Only what is needed for goal'],
     ['Best when','Many queries / derive all','Single specific goal'],
     ['Analogy','Data-driven (push)','Goal-driven (pull)'],
    ],widths=[1.4,2.7,2.5])

doc.add_page_break()

# ══════════════════════════════════════════════
h1('L04: Local Search')
# ══════════════════════════════════════════════

h3('Search Landscape Intuition')
eli5('Imagine a hilly landscape. Every point is a possible solution. Height = quality. You want to reach the highest peak.')
sp()

h3('Hill Climbing')
eli5('Always take a step uphill. Simple and fast, but you get stuck on small hills and can\'t see the big mountain.')
sp()
blue('Rule: ','At each step, move to the best neighbour. Stop when no neighbour is better.')
tbl(['Problem','Description','Example'],
    [['Local optimum','A hill that is not the global peak. HC cannot escape.','8-queen: stuck with 1 attack pair'],
     ['Plateau','Flat region — no neighbour is better or worse.','All neighbours equal score'],
     ['Ridge','Series of local optima in one direction','Diagonal sequence of hills'],
    ],widths=[1.4,2.8,2.2])
note('Exam tip: HC gets stuck ~50% of the time on 8-queens. SA and restarts are the fix.')

h3('Simulated Annealing (SA)')
eli5('Mostly go uphill, but sometimes randomly step DOWN. Early on (hot) you wander freely. Later (cool) you almost always go uphill. This lets you escape local optima.')
sp()
blue('Step 1: ','Start with initial state and temperature T.')
blue('Step 2: ','Pick a random neighbour.')
blue('Step 3: ','If better → always accept. If worse → accept with probability P = e^{-|delta|/T}.')
blue('Step 4: ','Cool T according to cooling schedule. Repeat until T ≈ 0.')
sp()
formula('P(accept worse move) = e^{-|delta_cost| / T}')
formula('|delta_cost| = |new_cost - old_cost|  (always positive for worsening move)')
sp()
tbl(['Temperature','Behaviour','Analogy'],
    [['T high (early)','P close to 1 — accept almost any move','Wandering phase, exploring freely'],
     ['T medium','P moderate — sometimes accept bad moves','Starting to focus'],
     ['T low (late)','P close to 0 — almost never accept bad','Converging phase'],
    ],widths=[1.5,2.8,2.1])
h4('SA Sign Convention — Three Forms, Same Answer')
tbl(['Convention','delta definition','Formula','Result for delta_cost=1, T=10'],
    [['A','delta = new - old (positive for worsening)','P = e^{-delta/T}','e^{-1/10}=0.905'],
     ['B','delta = old - new (negative for worsening)','P = e^{+delta/T}','e^{-1/10}=0.905'],
     ['C','delta = |new - old| (always positive)','P = e^{-delta/T}','e^{-1/10}=0.905'],
    ],widths=[0.8,2.4,1.8,2.4])
note('Exam tip: For a BETTER move, SA ALWAYS accepts — no probability needed. Write: "28 < 31, so SA accepts unconditionally."')

h3('Genetic Algorithm (GA)')
eli5('Have 100 possible answers. The good ones breed (swap parts). Random mutations flip bits. Bad ones die. Repeat — it\'s evolution.')
sp()
tbl(['Step','Name','Detail'],
    [['1','Representation','Each candidate = a chromosome (e.g. binary string). For feature selection: 1000-bit string, 1=selected feature.'],
     ['2','Population','N random chromosomes to start.'],
     ['3','Fitness','Measure how good each chromosome is. Feature selection: train SVM, fitness = validation accuracy.'],
     ['4','Selection','Higher fitness = more likely to be chosen as parent (roulette wheel or tournament).'],
     ['5','Crossover','Pick crossover point. Swap suffixes of two parents. Creates two children.'],
     ['6','Mutation','Randomly flip bits with probability Pm (~0.01). Maintains diversity.'],
     ['7','Termination','Stop at max generations or when fitness improvement < threshold.'],
    ],widths=[0.4,1.3,4.7])
note('Exam tip (Dec 2015 Q2.4 — 8 marks): GA for feature selection — be ready to describe all 7 steps with the binary encoding, SVM fitness, and 30k test evaluation at the end.')

doc.add_page_break()

# ══════════════════════════════════════════════
h1('L05: Constraint Satisfaction & Games')
# ══════════════════════════════════════════════

h3('CSP — Constraint Satisfaction Problem')
eli5('Colour a map so no two touching countries share a colour. You have rules (constraints) — figure out what values to assign each variable.')
sp()
bullet('Variables: things to assign (regions, course slots, etc.).')
bullet('Domains: possible values for each variable (colours, times, etc.).')
bullet('Constraints: rules that restrict which value combinations are legal.')
sp()
note('Exam tip: CSP solution = complete (all variables assigned) + consistent (no constraint violated).')

h3('Variable and Value Ordering Heuristics')
eli5('MCV asks "which country to colour NEXT?" LCV asks "which colour to USE for that country?" Degree is a tie-breaker.')
sp()
tbl(['Heuristic','Selects','Rule','Effect'],
    [['MCV (Most Constrained Variable)','Variable','Fewest remaining legal values','Detects dead-ends early'],
     ['LCV (Least Constraining Value)','Value','Fewest eliminations for neighbours','Keeps options open'],
     ['Degree','Variable','Most constraints (edges) with other unassigned vars','Good as MCV tie-breaker'],
    ],widths=[1.9,0.9,2.2,1.6])
note('Exam tip: MCV chooses the VARIABLE. LCV chooses the VALUE. Do not mix them up. Memory hook: Most Constrained Variable → Variable.')

h3('Minimax')
eli5('Two players take turns. MAX wants the highest score. MIN wants the lowest. Both play perfectly. Minimax figures out the best move for MAX assuming MIN plays optimally too.')
sp()
blue('Step 1: ','Build the full game tree to terminal states. Assign utility values to leaf nodes.')
blue('Step 2: ','Propagate values bottom-up: MIN nodes take minimum of children; MAX nodes take maximum.')
blue('Step 3: ','Root value = best outcome MAX can guarantee. Choose the action that leads to it.')
sp()
tbl(['Node type','Rule','Analogy'],
    [['MAX (your turn)','Take maximum of children\'s values','Pick best move for yourself'],
     ['MIN (opponent\'s turn)','Take minimum of children\'s values','Opponent picks worst for you'],
     ['Leaf (terminal)','Use utility function value directly','Game over — score counted'],
    ],widths=[1.6,2.8,2.0])
note('Exam tip: In a Nim/stone game, first identify who wins from each terminal state, then propagate.')

h3('Alpha-Beta Pruning')
eli5('Play Minimax but skip branches you ALREADY KNOW won\'t change the result. The opponent will never let you reach them anyway.')
sp()
blue('alpha: ','Best value MAX has found so far along current path. Starts at -infinity.')
blue('beta: ','Best value MIN has found so far along current path. Starts at +infinity.')
blue('Prune: ','At a MIN node, if current value <= alpha → prune (MAX already has better). At a MAX node, if current value >= beta → prune (MIN already has better).')
sp()
formula('Prune at MIN node: v <= alpha')
formula('Prune at MAX node: v >= beta')
sp()
note('Exam tip: Alpha-Beta produces IDENTICAL result to Minimax — just faster. Best case: prunes half the tree (O(b^{d/2}) instead of O(b^d)).')

doc.add_page_break()

# ══════════════════════════════════════════════
h1('L06: Machine Learning')
# ══════════════════════════════════════════════

h3('ML Taxonomy')
eli5('Supervised = teacher gives you correct answers. Unsupervised = no answers, find patterns yourself. Reinforcement = trial and error with rewards.')
sp()
tbl(['Type','Labels?','Goal','Examples'],
    [['Supervised','Yes','Learn input→output mapping','SVM, Decision Tree, MLP'],
     ['Unsupervised','No','Discover hidden structure','K-Means, PCA'],
     ['Reinforcement','Rewards','Maximise cumulative reward','AlphaGo, game agents'],
    ],widths=[1.5,0.7,2.3,2.1])

h3('K-Means Clustering')
eli5('Put K magnets on the data. Each point sticks to the nearest magnet. Move each magnet to the centre of its group. Repeat until the magnets stop moving.')
sp()
blue('Step 1: ','Choose K. Initialise K cluster centres (given in exam or random).')
blue('Step 2: ','Assign each point to the nearest centre by distance: |x - centre|.')
blue('Step 3: ','Recompute each centre as the mean of its assigned points.')
blue('Step 4: ','Repeat steps 2–3 until assignments do not change (convergence).')
sp()
h4('Dec 2023 Q4 — K-Means: 1D data, K=3 [12 marks — full iteration table required]')
blue('Data: ','0.12, 0.15, 0.21, 0.50, 0.69, 0.78, 0.89, 1.25, 1.46, 1.97')
blue('Initial centroids: ','μ1=0.12, μ2=0.78, μ3=1.97  (first, middle, last)')
tbl(['Point','|x-μ1=0.12|','|x-μ2=0.78|','|x-μ3=1.97|','Cluster Iter 1'],
    [['0.12','0.00','0.66','1.85','C1'],
     ['0.15','0.03','0.63','1.82','C1'],
     ['0.21','0.09','0.57','1.76','C1'],
     ['0.50','0.38','0.28','1.47','C2'],
     ['0.69','0.57','0.09','1.28','C2'],
     ['0.78','0.66','0.00','1.19','C2'],
     ['0.89','0.77','0.11','1.08','C2'],
     ['1.25','1.13','0.47','0.72','C3'],
     ['1.46','1.34','0.68','0.51','C3'],
     ['1.97','1.85','1.19','0.00','C3'],
    ],widths=[0.7,1.3,1.3,1.3,1.2])
formula('New centroids: μ1=(0.12+0.15+0.21)/3=0.16   μ2=(0.50+0.69+0.78+0.89)/4=0.715   μ3=(1.25+1.46+1.97)/3=1.56')
tbl(['Point','|x-μ1=0.16|','|x-μ2=0.715|','|x-μ3=1.56|','Cluster Iter 2'],
    [['0.12','0.04','0.595','1.44','C1'],
     ['0.15','0.01','0.565','1.41','C1'],
     ['0.21','0.05','0.505','1.35','C1'],
     ['0.50','0.34','0.215','1.06','C2'],
     ['0.69','0.53','0.025','0.87','C2'],
     ['0.78','0.62','0.065','0.78','C2'],
     ['0.89','0.73','0.175','0.67','C2'],
     ['1.25','1.09','0.535','0.31','C3'],
     ['1.46','1.30','0.745','0.10','C3'],
     ['1.97','1.81','1.255','0.41','C3'],
    ],widths=[0.7,1.3,1.35,1.35,1.2])
formula('Assignments identical to Iter 1 → CONVERGED in 2 iterations.')
formula('Final clusters: C1={0.12,0.15,0.21}  |  C2={0.50,0.69,0.78,0.89}  |  C3={1.25,1.46,1.97}')
note('Exam tip: Always show the full distance table for EACH iteration. Explicitly state "CONVERGED" when assignments stop changing. Partial marks require the table, not just the answer.')

h3('Decision Tree & Entropy')
eli5('A game of 20 questions. Pick the question that splits your group most cleanly. Entropy measures messiness — 0 = perfectly clean, 1.0 = 50/50 chaos.')
sp()
formula('H(S) = -sum_i p_i * log2(p_i)       [entropy of set S]')
formula('IG(F) = H(S) - sum_v (|S_v|/|S|) * H(S_v)   [information gain of feature F]')
sp()
blue('Step 1: ','Compute H(S) for the full dataset.')
blue('Step 2: ','For each feature F: compute the weighted entropy after splitting on F.')
blue('Step 3: ','IG(F) = H(S) - weighted entropy. Pick the feature with highest IG as the root.')
blue('Step 4: ','Recurse on each branch. Stop when entropy = 0 (pure) or no features left.')
sp()
tbl(['Composition','Entropy H'],
    [['All same class (e.g. 8Y/0N)','H = 0 (perfectly clean)'],
     ['50/50 (e.g. 4Y/4N)','H = 1.0 (maximally mixed)'],
     ['6Y/2N','H = -(6/8)*log2(6/8) - (2/8)*log2(2/8) = 0.811'],
    ],widths=[2.8,3.6])
note('Exam tip: Common exam pattern — 14 instances, 5 features. Always compute H(S) first, then IG for each feature. Root = highest IG.')

h4('Dec 2024 Q6 — Decision Tree (8 students, 3 features → Pass/Fail) [12 marks]')
tbl(['ID','Study Years','Study Hours','Attendance','Pass'],
    [['1','0-2','Low','I','No'],
     ['2','0-2','Low','R','No'],
     ['3','3-5','Medium','I','No'],
     ['4','3-5','Low','I','No'],
     ['5','3-5','High','R','Yes'],
     ['6','0-2','High','I','Yes'],
     ['7','0-2','High','R','Yes'],
     ['8','3-5','Medium','R','Yes'],
    ],widths=[0.4,1.0,1.0,1.1,0.6])
formula('H(S) = -4/8*log2(4/8) - 4/8*log2(4/8) = 1.0')
tbl(['Feature','Partitions (Y/N per group)','Partition H','IG'],
    [['Study Years','0-2:{1,2,6,7}=2Y2N; 3-5:{3,4,5,8}=2Y2N','1.0 each','1.0-(0.5*1.0+0.5*1.0)=0.0'],
     ['Study Hours','Low:{1,2,4}=0Y3N; High:{5,6,7}=3Y0N; Med:{3,8}=1Y1N','0; 0; 1.0','1.0-(3/8*0+3/8*0+2/8*1.0)=0.75'],
     ['Attendance','I:{1,3,4,6}=1Y3N; R:{2,5,7,8}=3Y1N','0.811 each','1.0-(0.5*0.811+0.5*0.811)=0.189'],
    ],widths=[1.1,2.4,1.3,1.8])
blue('Root: ','Study Hours (IG=0.75, highest)')
blue('Tree: ','Low → Predict No  |  High → Predict Yes  |  Medium → split on Attendance: I→No (ID3), R→Yes (ID8)')
formula('Study Years gives IG=0.0 — splitting on it conveys ZERO additional information. Never choose it as root or split.')
note('Exam tip: Study Years has equal class distribution in BOTH sub-groups (2Y/2N each) → pure chance → IG=0. This is a deliberate trap. Always compute IG for ALL features before deciding.')

h3('SVM — Support Vector Machine')
eli5('Draw the WIDEST possible street between two groups of dots. The dots right on the kerb are the support vectors. Maximise the street width.')
sp()
bullet('Support vectors: data points on the margin boundary — closest to the hyperplane.')
bullet('Decision boundary: w^T x + b = 0')
bullet('Margin = 2 / ||w||. Maximising margin = finding the best separator.')
bullet('Points on positive margin: w^T x + b = +1. Negative margin: w^T x + b = -1.')
sp()
h4('SVM Linear Classifier (Quiz 6 2025): SVs at (5,1) RED and (-1,-1) BLUE')
tbl(['Step','Calculation','Result'],
    [['Midpoint','((5-1)/2, (1-1)/2)','(2, 0)'],
     ['SV line slope','(1-(-1))/(5-(-1)) = 2/6','1/3'],
     ['Perpendicular slope','k_d * k_perp = -1  →  k_perp = -3','-3'],
     ['Intercept','0 = -3*2 + b  →  b = 6','6'],
     ['Decision boundary','y = -3x + 6','—'],
     ['Margin','2/||w|| = 2/sqrt(10)','0.632'],
    ],widths=[1.8,3.0,1.8])
note('Exam tip: Identify the two support vectors from the diagram (one per class), find the midpoint, then draw a line perpendicular to the line joining them.')

doc.add_page_break()

# ══════════════════════════════════════════════
h1('L07: Neural Networks')
# ══════════════════════════════════════════════

h3('Biological vs Artificial Neuron')
eli5('A real neuron has branches that receive signals (dendrites), a cell body that adds them all up, and a wire that sends the result forward (axon). The artificial version copies this exactly with numbers.')
sp()
tbl(['Biological','Artificial','Role'],
    [['Dendrites','Inputs x_i','Receive incoming signals'],
     ['Synaptic weights','Weights w_i','Strengthen or weaken each connection'],
     ['Soma (cell body)','Weighted sum: net_j = sum_i(w_i*x_i)','Integrate all signals'],
     ['Firing threshold','Bias weight w_0','Shift the activation threshold'],
     ['Axon','Output: o_j = sigma(net_j)','Transmit result to the next layer'],
    ],widths=[1.8,2.4,2.4])
blue('Scale: ','~10^11 neurons in the brain; each with ~10^4 synapses. Artificial nets: millions of weights.')

h3('Activation Functions')
eli5('The activation function is what makes the neuron "fire" or not. Sigmoid squishes everything smoothly into (0,1). Sign just shouts +1 or -1.')
sp()
tbl(['Function','Formula','Differentiable?','Use in backprop?'],
    [['Sign (step)','o = +1 if net>=0, -1 if net<0','No (zero derivative)','NO — gradient cannot flow'],
     ['Sigmoid','o = 1/(1+e^{-net})','Yes: sigma*(1-sigma)','YES — standard choice'],
     ['ReLU','o = max(0, net)','Yes (except at 0)','YES — common in deep nets'],
    ],widths=[1.3,2.2,1.5,1.6])
formula('sigma(x) = 1 / (1 + e^{-x})')
formula("sigma'(x) = sigma(x) * (1 - sigma(x))")
note('Exam tip: Sign function has zero derivative → zero gradient → weights cannot update → network cannot learn. This is the standard exam answer for why sigmoid is needed.')

h3('Perceptron Learning')
eli5('One neuron. Binary classification. If it guesses wrong, nudge every weight proportional to the mistake. If it guesses right, do nothing.')
sp()
formula('y_hat = sign(sum_i w_i * x_i)')
formula('w_new = w_old + eta * (y - y_hat) * x   [update only when y != y_hat]')
sp()
h4('Perceptron Example (Dec 2016 Q4.1) — 4D data, eta=0.1, w_0=(0.1,-0.02,-0.02,0.07)')
tbl(['Point','x','net = w^T x','y_hat','y','Update?'],
    [['x1','(0.14,-0.54,0.26,0.33)','0.014+0.011-0.005+0.023=0.043','+1','+1','No'],
     ['...','...','...','...','...','Continue for 8 points'],
    ],widths=[0.6,2.0,2.4,0.7,0.5,1.4])
note('Exam tip: Perceptron only converges if data is LINEARLY SEPARABLE. Always compute net = w^T x, take sign, compare to y, update if wrong.')

h3('MLP Backpropagation — Full Algorithm')
eli5('Stack layers of neurons. Forward: compute the output. Backward: measure how wrong it was, then blame each weight proportionally and nudge it.')
sp()
blue('FORWARD PASS:')
formula('net_j = sum_i W_ij * x_i          [hidden node j]')
formula('o_j   = sigma(net_j)')
formula('net_p = sum_j W_jp * o_j          [output node p]')
formula('y_hat_p = sigma(net_p)')
sp()
blue('OUTPUT DELTA:')
formula('delta_p = (y_p - y_hat_p) * y_hat_p * (1 - y_hat_p)')
sp()
blue('UPDATE OUTPUT WEIGHTS:')
formula('W_jp_new = W_jp + eta * delta_p * o_j')
sp()
blue('HIDDEN DELTA:')
formula('delta_j = o_j * (1 - o_j) * sum_p(W_jp * delta_p)')
sp()
blue('UPDATE HIDDEN WEIGHTS:')
formula('W_ij_new = W_ij + eta * delta_j * x_i')
sp()
h4('Dec 2023 Q7 — 4-input, 3-hidden, 2-output; all-ones weights; x=[1,0,1,0], y=[1,0], eta=0.5')
tbl(['Step','Calculation','Value'],
    [['net_j (all 3)','1*1+1*0+1*1+1*0 = 2','2 for j=1,2,3'],
     ['o_j (all 3)','sigma(2) = 1/(1+e^{-2})','0.8808'],
     ['net_p (both)','3 * 0.8808','2.6424'],
     ['y_hat (both)','sigma(2.6424)','0.9335'],
     ['delta_p1','(1-0.9335)*0.9335*0.0665','0.0041'],
     ['delta_p2','(0-0.9335)*0.9335*0.0665','-0.0579'],
     ['W(1)_jp1','1.0 + 0.5*0.0041*0.8808','1.0018 (all 3 rows)'],
     ['W(1)_jp2','1.0 + 0.5*(-0.0579)*0.8808','0.9745 (all 3 rows)'],
     ['sum_p(W_jp*delta_p)','1.0*0.0041 + 1.0*(-0.0579)','-0.0538'],
     ['delta_j (all 3)','0.8808*0.1192*(-0.0538)','-0.0057'],
     ['W(0) i=1,3 (x=1)','1.0 + 0.5*(-0.0057)*1','0.9971'],
     ['W(0) i=2,4 (x=0)','no change','1.0000'],
    ],widths=[2.0,3.0,1.6])
note('Exam tip: Final answer — W(1): all rows=[1.0018,0.9745]. W(0): all rows=[0.9971,1.0,0.9971,1.0].')

h3('Gradient Descent Variants')
tbl(['Variant','Updates after…','Properties'],
    [['Batch GD','All N examples','Stable, slow. One update per epoch.'],
     ['Stochastic GD','Each single example','Fast but noisy. Many updates per epoch.'],
     ['Mini-batch GD','Every B examples','Best of both. Most common in practice.'],
    ],widths=[1.5,1.8,3.3])
formula('Mini-batch: iterations per epoch = N / B')
formula('Total iterations = (N / B) * num_epochs')
note('Exam tip (May 2025 Q17b): 1000 samples, batch=50, 200 epochs → (1000/50)*200 = 4000 iterations.')

doc.add_page_break()

# ══════════════════════════════════════════════
h1('L08: Deep Learning')
# ══════════════════════════════════════════════

h3('CNN — Convolutional Neural Network')
eli5('For images. A small filter slides over the whole image. Each position: multiply every pixel in the filter area by the filter, add them up. Detects edges and textures locally. Then pooling shrinks the result so position doesn\'t matter. Finally a normal neural net gives the answer.')
sp()
tbl(['Layer','Function','Key property'],
    [['Convolutional','Slides filter (kernel) over input. Detects local features.','Shared weights — far fewer parameters than MLP'],
     ['Pooling (max/avg)','Reduces spatial dimensions. Takes max/mean in each region.','Translation invariance'],
     ['Fully Connected','Flattens all features. Classic MLP to produce output scores.','Same as standard MLP layer'],
    ],widths=[1.5,3.0,2.1])
sp()
formula('Output size = (image_size - kernel_size) / stride + 1')
formula('Example: 4x4 image, 3x3 kernel, stride=1  →  (4-3)/1+1 = 2  →  2x2 output')
note('Exam tip: CNN is suitable for ANY task with spatial structure: image classification, object detection, crowd counting, medical imaging.')

h3('CNN Convolution Calculation')
eli5('Slide the filter over one patch. Multiply each overlapping number pair, add them all up. That\'s one cell of the output.')
sp()
h4('May 2025 Q19b — kernel [[1,0,-1],[1,0,-1],[1,0,-1]] on patch [[4,3,2],[1,0,1],[-1,-2,-1]]')
formula('= 4*1+3*0+2*(-1) + 1*1+0*0+1*(-1) + (-1)*1+(-2)*0+(-1)*(-1)')
formula('= (4-2) + (1-1) + (-1+1) = 2 + 0 + 0 = 2')
sp()
h4('May 2025 Q19c — CNN for Crowd Counting: Do You Agree? [6 marks]')
blue('YES — CNN is well-suited for crowd counting. Three reasons:')
blue('1. Spatial Locality: ','Convolutional filters detect local patterns (heads, bodies) at every position via weight sharing — critical when crowd density varies across the image.')
blue('2. Translation Invariance: ','Pooling makes detection robust to position shifts. A head at top-left and one at bottom-right produce the same filter response.')
blue('3. Density Map Output: ','CNNs can output a density map (one value per region) rather than a single count, enabling spatial localisation of crowd hotspots.')
blue('Enhancement techniques: ','(a) Multi-column CNN (MCNN) — parallel branches with different kernel sizes handle scale variation. (b) Dilated convolutions increase receptive field without losing resolution. (c) Attention mechanisms focus on head/body regions.')
note('6-mark rubric: ~2 marks per well-explained CNN property + 1 enhancement. Name the property AND explain why it helps counting specifically.')
sp()
h4('Aug 2022 Q4.2 — all-ones 3x3 kernel on 4x4 image')
blue('Rule: ','All-ones kernel = sum of the 9 pixels in the 3x3 region.')
formula('Output (2x2): [[1140, 1150], [1200, 1210]]')
note('Exam tip: For all-ones kernel, just sum the 9 pixels in each 3x3 region. For other kernels, multiply element-wise then sum.')

h3('RNN — Recurrent Neural Network')
eli5('Reads a sentence one word at a time. Has a hidden state — memory of what it just read. The problem: the memory fades quickly. Can\'t remember something from 100 words ago.')
sp()
bullet('Input at each step: current token + previous hidden state.')
bullet('Output: new hidden state (passed to next step) + optional output prediction.')
bullet('Problem: vanishing gradient — signals from early in the sequence shrink to zero during backprop.')

h3('LSTM — Long Short-Term Memory')
eli5('An RNN with a separate long-term memory lane (cell state) and three gates that control what to remember or erase.')
sp()
tbl(['Gate','Controls','Plain English'],
    [['Forget gate','What to erase from cell state','\"I can forget old context\"'],
     ['Input gate','What new information to write in','\"I should remember this new word\"'],
     ['Output gate','What part of memory to expose as output','\"Here\'s what\'s relevant right now\"'],
    ],widths=[1.5,2.5,2.6])
note('Exam tip: LSTM vs RNN — LSTM has cell state (long-term memory) + 3 gates. RNN has only hidden state. LSTM solves vanishing gradient.')

h3('GAN — Generative Adversarial Network')
eli5('Two networks fight each other. The forger (Generator) tries to fake images so good the detective (Discriminator) can\'t tell. The detective gets better. The forger gets better. Arms race. Eventually the forger is perfect.')
sp()
tbl(['Network','Role','Wins when…'],
    [['Generator G','Creates fake samples from random noise, tries to fool D','D outputs 0.5 (can\'t distinguish)'],
     ['Discriminator D','Classifies real vs fake, tries to catch G','Correctly identifies all fakes'],
    ],widths=[1.6,3.0,2.0])
blue('Equilibrium: ','G perfectly replicates the real data distribution. D outputs 0.5 for everything (random guess).')
note('Exam tip: GAN = generative model. G never sees real data directly — it only gets D\'s feedback (real/fake). D sees both real data and G\'s fakes.')

doc.add_page_break()

# ══════════════════════════════════════════════
h1('L09: Bayesian Networks & Naive Bayes')
# ══════════════════════════════════════════════

h3('Probability & Bayes Rule')
eli5('You have a belief about something. You see evidence. Bayes rule tells you how to update your belief based on that evidence.')
sp()
tbl(['Term','Meaning','Symbol'],
    [['Prior','Your belief BEFORE seeing evidence','P(H)'],
     ['Likelihood','How probable is this evidence IF the hypothesis is true?','P(E|H)'],
     ['Posterior','Your belief AFTER seeing evidence — what you want','P(H|E)'],
     ['Marginal likelihood','Total probability of evidence (normaliser)','P(E)'],
    ],widths=[1.5,3.2,1.9])
sp()
formula('Bayes Rule: P(H|E) = P(E|H) * P(H) / P(E)')
formula('P(E) = sum_h P(E|H=h) * P(H=h)   [total probability — sum over all hypotheses]')
sp()

h3('Bayesian Network Inference')
eli5('A graph of cause-and-effect arrows. Each node has a probability table. You can ask: given that I see THIS, what is the probability of THAT?')
sp()
bullet('Predictive inference (cause → effect): P(Wet|Rain). Forward in the graph.')
bullet('Diagnostic inference (effect → cause): P(Rain|Wet). Backward — use Bayes rule.')
bullet('Key calculation: sum over all parent combinations. Multiply joint probabilities.')
sp()
h4('Dec 2023 Q5 — Network: E→S, E→J, S→C, J→C  [P(C|E=T) and P(E|C)]')
tbl(['S','J','P(S,J|E=T)','P(C|S,J)','Contribution to P(C|E=T)'],
    [['S','J','0.2*0.8=0.160','0.9','0.144'],
     ['S','~J','0.2*0.2=0.040','0.4','0.016'],
     ['~S','J','0.8*0.8=0.640','0.8','0.512'],
     ['~S','~J','0.8*0.2=0.160','0.1','0.016'],
    ],widths=[0.5,0.5,1.8,1.3,2.5])
formula('P(C|E=T) = 0.144+0.016+0.512+0.016 = 0.688')
formula('P(C|E=F) = 0.338 (computed similarly with P(S|~E)=0.6, P(J|~E)=0.1)')
formula('P(C) = 0.688*0.3 + 0.338*0.7 = 0.4430')
formula('P(E|C) = P(C|E)*P(E)/P(C) = 0.688*0.3/0.4430 = 0.4659')
sp()
h4('Dec 2024 Q5 — Network: W→T, W→R, T→C, R→C  [P(C) and P(W|C)]')
formula('P(C|W=T) = 0.7*0.8*0.9+0.7*0.2*0.6+0.3*0.8*0.7+0.3*0.2*0.1 = 0.762')
formula('P(C|W=F) = 0.2*0.1*0.9+0.2*0.9*0.6+0.8*0.1*0.7+0.8*0.9*0.1 = 0.254')
formula('P(C) = 0.762*0.6 + 0.254*0.4 = 0.5588')
formula('P(W|C) = P(C|W)*P(W)/P(C) = 0.762*0.6/0.5588 = 0.8182')
note('Exam tip: Always enumerate ALL parent combinations systematically. A table with columns [parent combo | P(parents|evidence) | P(child|parents) | contribution] never misses a term.')

h3('Naive Bayes Classifier')
eli5('Classify something by multiplying: how common is this class? Times how likely is each feature given this class? The class with the biggest product wins. Naive = we pretend all features are independent.')
sp()
formula('Score(C) = P(C) * P(f1|C) * P(f2|C) * ... * P(fn|C)')
formula('Predicted class = argmax_C Score(C)')
sp()
h4('May 2025 Q18 — Full question has TWO parts')
blue('Part a) [8 marks] — ','Key assumption + DAG diagram')
blue('Key assumption: ','Naive Bayes assumes CONDITIONAL INDEPENDENCE — given the class label (Risk), each feature (Income, Credit History, Late Payment) is independent of all other features.')
formula('P(Income, Credit, Late | Risk) = P(Income|Risk) × P(Credit|Risk) × P(Late|Risk)')
blue('DAG diagram: ','Risk is the single parent node. Arrows point FROM Risk TO each feature.')
formula('         Risk')
formula('        /  |  \\')
formula('       ↓   ↓   ↓')
formula('  Income Credit Late')
formula('        History Payment')
note('Exam tip: In the Naive Bayes DAG, the CLASS is always the parent. All features are children with arrows pointing FROM class TO features. This is opposite to a causal model.')
sp()
blue('Part b) [12 marks] — ','Full classification calculation')
blue('Step 1 — Count training data:')
tbl(['Class','Count','IDs','Prior'],
    [['High Risk (HR)','3','1 (Low,Bad,Yes), 3 (High,Bad,Yes), 7 (High,Good,No)','P(HR) = 3/7'],
     ['Low Risk (LR)','4','2 (Low,Bad,No), 4 (High,Good,No), 5 (Low,Good,No), 6 (High,Good,Yes)','P(LR) = 4/7'],
    ],widths=[1.3,0.6,3.4,1.1])
blue('Step 2 — Conditional probability tables:')
tbl(['Feature','Value','HR count (of 3)','P(val|HR)','LR count (of 4)','P(val|LR)'],
    [['Income','High','ID3,ID7 = 2','2/3','ID4,ID6 = 2','2/4'],
     ['Credit History','Bad','ID1,ID3 = 2','2/3','ID2 = 1','1/4'],
     ['Late Payment','No','ID7 = 1','1/3','ID2,ID4,ID5 = 3','3/4'],
    ],widths=[1.3,0.7,1.5,0.9,1.5,0.9])
note('COMMON ERROR: P(Income=High|LR) = 2/4 NOT 3/4. LR instances with High income: only ID4 and ID6 = 2 out of 4.')
blue('Step 3 — Score calculation (Income=High, Credit=Bad, Late=No):')
formula('Score(HR) = 3/7 × 2/3 × 2/3 × 1/3 = 12/189 = 0.0635')
formula('Score(LR) = 4/7 × 2/4 × 1/4 × 3/4 = 24/448 = 0.0536')
formula('Score(HR) = 0.0635 > Score(LR) = 0.0536  →  Classified as HIGH RISK')
note('Exam tip: You do NOT need to normalise the scores — just compare them. The class with the highest product wins. This customer is HIGH RISK despite having High income, because Bad credit history has very low P(Bad|LR)=1/4.')

doc.add_page_break()

# ══════════════════════════════════════════════
h1('Quick Reference Summary')
# ══════════════════════════════════════════════

h3('Must-Know Formulas')
tbl(['Formula','Expression','When to use'],
    [['Sigmoid','sigma(x) = 1/(1+e^{-x})','MLP forward pass, perceptron output'],
     ["Sigmoid derivative","sigma'(x) = sigma(x)*(1-sigma(x))",'Backprop delta calculation'],
     ['MLP output delta','delta_p = (y_p - y_hat_p)*y_hat_p*(1-y_hat_p)','Output layer backprop'],
     ['MLP hidden delta','delta_j = o_j*(1-o_j)*sum_p(W_jp*delta_p)','Hidden layer backprop'],
     ['Weight update','W_new = W_old + eta*delta*input','Both layers in backprop'],
     ['A* evaluation','f(n) = g(n) + h(n)','Priority in A* open list'],
     ['SA acceptance','P = e^{-|delta_cost|/T}','Accepting worse moves (minimisation)'],
     ['Entropy','H(S) = -sum p*log2(p)','Decision tree node purity'],
     ['Information Gain','IG(F) = H(S) - sum(|Sv|/|S|)*H(Sv)','Choosing split feature in DT'],
     ['Bayes Rule','P(H|E) = P(E|H)*P(H)/P(E)','All Bayesian inference questions'],
     ['Naive Bayes','Score = P(C)*prod P(fi|C)','Classification score (pick argmax)'],
     ['PAC complexity','n >= (1/eps)*ln(|H|/delta)','PAC learning theory'],
     ['CNN output size','(img_size - kernel_size)/stride + 1','Checking output dimensions'],
     ['Mini-batch iters','Total = (N/B)*epochs','May 2025 Q17b type question'],
    ],widths=[1.8,2.6,2.2])

h3('Key Comparisons at a Glance')
tbl(['Pair','Key difference'],
    [['BFS vs UCS','BFS: fewest hops. UCS: cheapest cost. Same if all edge costs equal.'],
     ['UCS vs A*','A* adds heuristic h(n) to guide search. Same as UCS if h=0.'],
     ['Greedy vs A*','Greedy: f=h only (ignores past). A*: f=g+h (optimal). Greedy is faster but not optimal.'],
     ['FC vs BC','FC: data-driven, derives all consequences. BC: goal-driven, proves specific goal.'],
     ['MCV vs LCV','MCV picks next VARIABLE (fewest remaining values). LCV picks VALUE for that variable.'],
     ['HC vs SA','HC: never accepts worse. SA: sometimes accepts worse (escapes local optima).'],
     ['Discriminative vs Generative','Discriminative: P(Y|X). Generative: P(X,Y). Discriminative classifies; generative creates.'],
     ['Classification vs Clustering','Classification: supervised, predefined labels. Clustering: unsupervised, finds natural groups.'],
     ['Sign vs Sigmoid','Sign: not differentiable → no backprop. Sigmoid: differentiable everywhere → backprop works.'],
     ['RNN vs LSTM','RNN: hidden state only, vanishing gradient. LSTM: cell state + 3 gates, long-range memory.'],
    ],widths=[2.0,4.6])

h3('30-Second Calculation Recipes')
tbl(['Calculation','One-liner recipe'],
    [['BFS','FIFO queue. Alphabetical neighbours. Ignore costs. Fewest hops.'],
     ['DFS','Stack (go deepest). Alphabetical. Not optimal.'],
     ['UCS','Priority queue by g(n). Expand cheapest. Update if cheaper path found.'],
     ['A*','Priority queue by f=g+h. Expand lowest f. Update node if cheaper path found.'],
     ['K-Means','Assign to nearest centre → recompute centres as means → repeat until stable.'],
     ['Decision Tree','H(S) → IG for each feature → pick highest IG as root → recurse each branch.'],
     ['Perceptron','For each point: if wrong, w += eta*(y-y_hat)*x. If correct, do nothing.'],
     ['MLP Backprop','Forward → delta_out=(y-y_hat)*y_hat*(1-y_hat) → W(1) update → delta_hidden=o*(1-o)*sum(W*delta) → W(0) update.'],
     ['Naive Bayes','Score = prior * product of likelihoods per feature. Pick class with highest score.'],
     ['Bayesian net','Enumerate all parent combos. Sum: P(parents|evidence)*P(child|parents). Use Bayes for diagnostic.'],
     ['SA acceptance','Better → always accept. Worse → P = e^{-|new-old|/T}. State clearly which case.'],
     ['GD single neuron','dE/dwi = (A-Y)*A*(1-A)*xi. Then w_new = w - rate*dE/dw for each weight and bias.'],
     ['CNN convolution','Element-wise multiply kernel with patch, sum all products. Output size = (img-ker)/stride+1.'],
     ['Alpha-Beta prune','Track alpha (MAX best) and beta (MIN best). Prune when alpha>=beta.'],
    ],widths=[1.8,4.8])

h3('Exam Strategy')
tbl(['Question type','Strategy'],
    [['Short paragraph (2-4 marks)','Define the concept, give one example, state key advantage/limitation. 3-4 sentences max.'],
     ['Comparison (4-6 marks)','One sentence per difference. Use a mini-table if time allows. End with a use-case.'],
     ['Calculation — search (8-12 marks)','Draw the expansion table step by step. Show ALL frontier nodes with f/g/h values. State path at the end.'],
     ['Calculation — backprop (8-12 marks)','Forward pass first (show net_j and o_j). Output deltas. Update W(1). Hidden deltas. Update W(0).'],
     ['Calculation — Bayesian (8-12 marks)','Draw the enumeration table. Label columns clearly. Sum contributions. Apply Bayes for diagnostic.'],
     ['MCQ','Eliminate obviously wrong answers first. Watch for: DL is Connectionism not Symbolicism; LSTM not CNN for sequences.'],
    ],widths=[2.2,4.4])

doc.save('COMP4431_Full_Exam_Notes.docx')
print('Saved COMP4431_Full_Exam_Notes.docx')

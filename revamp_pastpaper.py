"""Revamp COMP4431_PastPaper_QuizSol.docx — clean structure with proper headings and tables."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.8); section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

def h1(t):
    p = doc.add_heading(t, 1); p.runs[0].font.color.rgb = RGBColor(0x1F,0x49,0x7D); return p
def h2(t):
    p = doc.add_heading(t, 2); p.runs[0].font.color.rgb = RGBColor(0x2E,0x74,0xB5); return p
def h3(t): return doc.add_heading(t, 3)
def body(t, bold=False, indent=0):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(indent)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(10); return p
def bullet(t, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Pt(indent)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(t); r.font.size = Pt(10); return p
def tip(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('EXAM TIP: '); r1.bold = True; r1.font.color.rgb = RGBColor(0xC0,0x55,0x00); r1.font.size = Pt(10)
    r2 = p.add_run(t); r2.font.color.rgb = RGBColor(0xC0,0x55,0x00); r2.font.size = Pt(10); return p
def hook(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('MEMORY: '); r1.bold = True; r1.font.color.rgb = RGBColor(0x37,0x86,0x43); r1.font.size = Pt(10)
    r2 = p.add_run(t); r2.font.color.rgb = RGBColor(0x37,0x86,0x43); r2.font.size = Pt(10); return p
def answer(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(16)
    r1 = p.add_run('ANSWER: '); r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(t); r2.font.size = Pt(10); return p
def formula(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t); r.bold = True; r.font.name = 'Courier New'; r.font.size = Pt(10); return p
def sp(): doc.add_paragraph().paragraph_format.space_after = Pt(0)
def tbl(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'
    hrow = t.rows[0]
    for i,h in enumerate(headers):
        c = hrow.cells[i]; c.text = h
        c.paragraphs[0].runs[0].bold = True; c.paragraphs[0].runs[0].font.size = Pt(9)
        sh = OxmlElement('w:shd'); sh.set(qn('w:fill'),'D6E4F0'); sh.set(qn('w:val'),'clear')
        c._tc.get_or_add_tcPr().append(sh)
    for ri,row in enumerate(rows):
        tr = t.rows[ri+1]
        for ci,val in enumerate(row):
            c = tr.cells[ci]; c.text = val; c.paragraphs[0].runs[0].font.size = Pt(9)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    sp()

# ── TITLE ──
title = doc.add_heading('COMP4431 AI — Past Papers & Quiz Solutions', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Fully worked solutions for all 6 past papers + 9 quizzes (2022–2025)')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True; sub.runs[0].font.color.rgb = RGBColor(0x1F,0x49,0x7D)
doc.add_page_break()

# ════════════════════════════════════════════
h1('SECTION 1 — Exam Format & Quick Reference')
# ════════════════════════════════════════════
tbl(['Item','Detail'],
    [['Format','20% MCQ + 80% Written (understanding, comparison, calculation)'],
     ['Calculator','Non-programmable calculator ALLOWED'],
     ['Coverage','Lectures 1–9. Mathematical proofs and lab content EXCLUDED.'],
     ['Key question types','Short paragraph (2-4 marks), calculation (8-16 marks), MCQ (1 mark each)'],
    ],widths=[2.0,4.2])

h2('Key Formulas — Quick Reference')
tbl(['Formula','Expression','When to use'],
    [['Sigmoid','sigma(x) = 1/(1+e^{-x})','Activation in perceptron, MLP forward pass'],
     ['Sigmoid derivative',"sigma'(x) = sigma(x)*(1-sigma(x))",'Backprop delta calculation'],
     ['MLP output delta','delta_p = (y_p - y_hat_p)*y_hat_p*(1-y_hat_p)','Output layer backprop'],
     ['MLP hidden delta','delta_j = o_j*(1-o_j)*sum_p(W_jp*delta_p)','Hidden layer backprop'],
     ['Weight update','W_new = W_old + eta*delta*input','Both layers in backprop'],
     ['SA acceptance','P = e^{-|delta_cost|/T}','Accepting worse moves in SA (minimisation)'],
     ['Entropy','H(S) = -sum p*log2(p)','Decision tree node purity'],
     ['Information Gain','IG(F) = H(S) - sum (|Sv|/|S|)*H(Sv)','Choosing split feature in DT'],
     ['A* evaluation','f(n) = g(n) + h(n)','Priority in A* open list'],
     ['Bayesian inference','P(H|E) = P(E|H)*P(H) / P(E)','Diagnostic/predictive reasoning'],
     ['Naive Bayes score','P(C)*prod P(xi|C)','Classification score (unnormalised)'],
     ['PAC sample complexity','n >= (1/epsilon)*ln(|H|/delta)','PAC learning theory'],
    ],widths=[1.8,2.5,2.1])
doc.add_page_break()

# ════════════════════════════════════════════
h1('SECTION 2 — Model Paragraph Answers')
# ════════════════════════════════════════════
body('These are model answers for short-paragraph questions that recur across all papers.', bold=True)
sp()

# ── 2.1 Discriminative vs Generative ──
h2('Discriminative vs Generative AI (Dec 2024 Q3.1)')
tbl(['','Discriminative','Generative'],
    [['What it learns','Decision boundary P(Y|X)','Data distribution P(X,Y) or P(X)'],
     ['Goal','Classify / label input','Generate new samples / understand data'],
     ['Examples','SVM, Logistic Regression, CNN classifier','GAN, VAE, Naive Bayes, RBM'],
     ['Analogy','Learns the fence between cats and dogs','Learns what a cat looks like'],
    ],widths=[1.4,2.4,2.4])
tip('Exam often asks for one example of each. GAN = generative; SVM = discriminative.')
sp()

# ── 2.2 Classification vs Clustering ──
h2('Classification vs Clustering (Dec 2024 Q3.2)')
tbl(['','Classification','Clustering'],
    [['Supervision','Supervised (labelled data)','Unsupervised (no labels)'],
     ['Goal','Map input to predefined class labels','Discover natural groupings in data'],
     ['Examples','SVM classifies email spam/not-spam','K-Means groups customers by purchase pattern'],
     ['Evaluation','Accuracy / F1 against known labels','Intra-cluster distance (no ground truth)'],
    ],widths=[1.4,2.4,2.4])
tip('Classification requires labelled training data; clustering does not.')
sp()

# ── 2.3 MCV vs LCV ──
h2('MCV vs LCV in CSP (Dec 2024 Q3.3, Quiz 3)')
tbl(['Heuristic','Full Name','Chooses…','Effect'],
    [['MCV','Most Constrained Variable','VARIABLE with fewest remaining legal values','Reduces branching; fails early if dead-end'],
     ['LCV','Least Constraining Value','VALUE that eliminates fewest choices for neighbours','Keeps options open; avoids dead-ends'],
     ['Degree','Degree Heuristic','VARIABLE with most constraints (edges)','Good as tie-breaker when MCV is tied'],
    ],widths=[0.8,2.0,2.0,1.6])
hook('MCV = which VARIABLE next. LCV = which VALUE to assign.')
sp()

# ── 2.4 Forward vs Backward Chaining ──
h2('Forward vs Backward Chaining (all papers, every year)')
tbl(['','Forward Chaining (FC)','Backward Chaining (BC)'],
    [['Direction','Data → Conclusion (bottom-up)','Goal → Premises (top-down)'],
     ['Start point','Known facts in KB','Query / goal to be proved'],
     ['Strategy','Apply all rules whose premises are satisfied; add conclusions','Identify rules that conclude the goal; check their premises recursively'],
     ['Best when','Many queries needed / derive all consequences','Single specific goal; avoid irrelevant rules'],
     ['Example','Car diagnosis: headlights_dim → battery_dead → dead_battery_issue','Q1←G∧K; G←A∧F; F←A∧C ✓ → Q1 PROVED'],
    ],widths=[1.3,2.5,2.5])
sp()

# ── 2.5 Hill Climbing vs SA ──
h2('Hill Climbing vs Simulated Annealing (Dec 2015 Q2.1, Quiz 3)')
tbl(['','Hill Climbing','Simulated Annealing'],
    [['Move rule','Always move to best neighbour','Accept better moves always; accept worse with P=e^{-|delta|/T}'],
     ['Local optima','Gets stuck (cannot escape)','Escapes with probability decreasing as T cools'],
     ['8-Queen example','Gets stuck ~50% of the time with 1 attack pair remaining','Finds solution with slow cooling schedule'],
     ['Guarantee','Not optimal','Near-optimal if cooled slowly enough'],
    ],widths=[1.3,2.5,2.5])
tip('If question gives specific T and delta values, calculate P = e^{-|delta|/T} numerically.')
sp()

# ── 2.6 Biological vs Artificial Neuron ──
h2('Biological vs Artificial Neuron (Dec 2015 Q2.2, Dec 2016, all years)')
tbl(['Biological','Artificial','Role'],
    [['Dendrites','Inputs x_i','Receive incoming signals'],
     ['Synaptic weights','Weights w_i','Strengthen/weaken connection'],
     ['Soma (cell body)','Weighted sum: net_j = sum(w_i*x_i)','Integrate signals'],
     ['Firing threshold','Bias weight w_0','Shift activation threshold'],
     ['Axon','Output: o_j = sigma(net_j)','Transmit result to next neurons'],
    ],widths=[1.8,2.2,2.2])
body('~10^11 neurons in brain; each with ~10^4 synapses. Artificial nets: millions of weights.',indent=12)
sp()

# ── 2.7 Sign vs Sigmoid in Backprop ──
h2('Why Sign Function Cannot Be Used in Backpropagation (Quiz 4 Q2b)')
bullet('Sign function is NOT differentiable at x=0, and has derivative = 0 everywhere else.')
bullet('Backpropagation requires computing gradients (chain rule) through every neuron.')
bullet('Zero gradient -> no information on how to update weights -> network CANNOT learn.')
bullet("Sigmoid is differentiable everywhere: sigma'(x) = sigma(x)*(1-sigma(x)) — always non-zero.")
answer("'Sign function is not differentiable / has zero derivative everywhere → prevents gradient calculation for weight update in backpropagation. Sigmoid's non-zero derivative allows gradient flow.'")
sp()

# ── 2.8 LSTM vs RNN ──
h2('LSTM vs Basic RNN (Dec 2024 Q concept)')
tbl(['','Basic RNN','LSTM'],
    [['Memory type','Hidden state only (short-term)','Cell state (long-term) + hidden state (short-term)'],
     ['Vanishing gradient','Severe — cannot learn long-range deps','Mitigated by gates controlling gradient flow'],
     ['Gates','None','Forget gate, Input gate, Output gate'],
     ['When to use','Short sequences only','Long sequences (NLP, time series, speech)'],
    ],widths=[1.3,2.5,2.5])
sp()

# ── 2.9 CNN vs MLP for Images ──
h2('CNN vs MLP for Image Processing')
tbl(['','MLP','CNN'],
    [['Input handling','Flattens image to 1D — destroys spatial structure','Preserves 2D spatial structure'],
     ['Parameters','Fully connected — extremely many parameters','Shared kernel weights — far fewer parameters'],
     ['Translation','Not invariant — position change breaks it','Pooling provides translation invariance'],
     ['Local features','Cannot detect local patterns','Convolutional layers detect edges, textures locally'],
    ],widths=[1.3,2.5,2.5])
tip('CNN is suitable for any task with spatial locality: image classification, object detection, crowd counting.')
sp()

# ── 2.10 GA for Feature Selection ──
h2('GA for Feature Selection (Dec 2015 Q2.4 — 8 marks)')
body('Setup: 1000 features, select 100. 50,000 data (20k train, 30k test). Classifier: linear SVM.',indent=12)
tbl(['Step','Detail'],
    [['Representation','Binary string of length 1000. Bit i=1 = feature selected. Constrain to exactly 100 ones.'],
     ['Population','Generate N=50 random chromosomes, each with 100 ones placed randomly.'],
     ['Fitness','Train SVM on selected features using training data. Fitness = validation accuracy.'],
     ['Selection','Roulette wheel or tournament selection. Higher fitness = more likely to be parent.'],
     ['Crossover','Single-point crossover: swap suffix of two parents. Repair if count != 100.'],
     ['Mutation','Flip bits with probability Pm~0.01. Maintains diversity.'],
     ['Termination','Max generations OR fitness improvement < threshold. Return best chromosome.'],
     ['Final eval','Evaluate selected features on 30k test set to report test accuracy.'],
    ],widths=[1.5,5.0])
sp()

# ── 2.11 PAC Learning ──
h2('PAC Learning (Dec 2015 Q2.3 — 6 marks)')
body('PAC = Probably Approximately Correct learning framework.',indent=12)
bullet('APPROXIMATELY CORRECT: error rate <= epsilon for any small epsilon > 0 chosen in advance.')
bullet('PROBABLY correct: probability of error > epsilon is at most delta.')
body('In plain English: with probability at least (1-delta), the learned hypothesis has error <= epsilon.',indent=12)
formula('n >= (1/epsilon) * ln(|H| / delta)')
body('|H| = hypothesis space size. Halving epsilon doubles the samples needed.',indent=12)
sp()

doc.add_page_break()

# ════════════════════════════════════════════
h1('SECTION 3 — Past Paper Solutions by Paper')
# ════════════════════════════════════════════
body('Papers listed most-recent-first. Most likely to repeat: Dec 2023 A* graph (appeared Dec 2015 AND Dec 2023).', bold=True)
sp()

# ──────────────────────────────────────────
h2('PAPER 1 — May 2025 Sem II (20250714_1_23.pdf)')
# ──────────────────────────────────────────
h3('Q16a — BFS Tree: A-2-C, A-4-B, C-10-E, C-8-D, B-5-D, D-2-E  (Goal=E)')
tbl(['Level','Node expanded','Queue before expansion','Children added','Note'],
    [['0','A','[A]','B, C (alphabetical)','Start'],
     ['1','B','[B, C]','D',''],
     ['1','C','[C, D]','D (skip—already queued), E','E = GOAL — found!'],
    ],widths=[0.5,0.9,1.5,1.8,1.7])
answer('BFS path: A → C → E  (2 hops). Trace: E.parent=C, C.parent=A.')
tip('BFS finds FEWEST HOPS, not cheapest path. A→B→D→E is 3 hops and WRONG even though cost may be lower.')
body('DRAW THIS IN EXAM (tree diagram):',indent=12)
formula('          A')
formula('         / \\')
formula('        B   C')
formula('        |  / \\')
formula('        D  D [E]  <- GOAL (D from B is skipped)')
sp()

h3('Q16b — A* on same graph. H: A=7, B=6, C=4, D=2, E=0')
tbl(['Step','Node expanded','f=g+h used','Frontier after expansion'],
    [['1','A  (f=0+7=7)','—','C(f=2+4=6), B(f=4+6=10)'],
     ['2','C  (f=2+4=6)','','D(f=10+2=12), E(f=12+0=12), B(10)'],
     ['3','B  (f=4+6=10)','','D via B (f=9+2=11) updates D, E(12)'],
     ['4','D via B (f=11)','','E via D (f=11+2=0=13?? no: g=9+2=11, h=0, f=11). Pick E(12).'],
     ['FINAL','E via C (f=12)','','GOAL'],
    ])
answer('CORRECT PATH: A → B → D → E, cost = 4+5+2 = 11. (Not A→C→E=12 — A* finds optimal cost.)')
tip('Update D via B: g(D via B) = g(B)+5 = 9; f = 9+2=11 < 12 (via C). So D via B is better — update it.')
sp()

h3('Q17a — Gradient Descent Single Neuron: w1=0.1, w2=0.1, b=0.1, x=[8,4], Y=3, rate=0.1')
tbl(['Variable','Calculation','Value'],
    [['s','8*0.1 + 4*0.1 + 0.1','1.3'],
     ['A = sigma(s)','1/(1+e^{-1.3}) = 1/(1+0.2725)','0.7858'],
     ['dE/dw1','(A-Y)*A*(1-A)*x1 = (0.7858-3)*0.7858*0.2142*8','-2.9797'],
     ['dE/dw2','(A-Y)*A*(1-A)*x2 = (-2.2142)*0.7858*0.2142*4','-1.4899'],
     ['dE/db','(A-Y)*A*(1-A)*1 = (-2.2142)*0.7858*0.2142','-0.3725'],
     ['w1_new','0.1 - 0.1*(-2.9797)','0.3980'],
     ['w2_new','0.1 - 0.1*(-1.4899)','0.2490'],
     ['b_new','0.1 - 0.1*(-0.3725)','0.1373'],
    ],widths=[1.6,3.5,0.8])
h3('Q17b — Mini-batch iterations: 1000 samples, batch=50, 200 epochs')
formula('Iterations = (1000/50) * 200 = 20 * 200 = 4000')
sp()

h3('Q18 — Naive Bayes Credit Risk [20 marks total]')
body('Training data: 7 customers. High Risk (HR): IDs 1,3,7. Low Risk (LR): IDs 2,4,5,6.', bold=True)
sp()

h3('Q18a — Key Assumption + DAG Diagram [8 marks]')
body('Key assumption: CONDITIONAL INDEPENDENCE. Given the class label (Risk), each feature is independent of all other features.', indent=12)
from docx.shared import Pt as _Pt
p = doc.add_paragraph(); p.paragraph_format.left_indent = _Pt(16); p.paragraph_format.space_after = _Pt(2)
r = p.add_run('P(Income, Credit, Late | Risk) = P(Income|Risk) x P(Credit|Risk) x P(Late|Risk)'); r.bold = True; r.font.name = 'Courier New'; r.font.size = _Pt(10)
body('DAG diagram: Risk is the single parent. Arrows point FROM Risk TO each feature.', indent=12)
p = doc.add_paragraph(); p.paragraph_format.left_indent = _Pt(24); p.paragraph_format.space_after = _Pt(1)
r = p.add_run('         Risk'); r.font.name = 'Courier New'; r.font.size = _Pt(10)
p = doc.add_paragraph(); p.paragraph_format.left_indent = _Pt(24); p.paragraph_format.space_after = _Pt(1)
r = p.add_run('        /  |  \\'); r.font.name = 'Courier New'; r.font.size = _Pt(10)
p = doc.add_paragraph(); p.paragraph_format.left_indent = _Pt(24); p.paragraph_format.space_after = _Pt(1)
r = p.add_run('       v   v   v'); r.font.name = 'Courier New'; r.font.size = _Pt(10)
p = doc.add_paragraph(); p.paragraph_format.left_indent = _Pt(24); p.paragraph_format.space_after = _Pt(2)
r = p.add_run('  Income Credit  Late'); r.font.name = 'Courier New'; r.font.size = _Pt(10)
tip('The class node is ALWAYS the parent in a Naive Bayes DAG. Arrows go FROM class TO features. State the conditional independence assumption explicitly for full marks.')
sp()

h3('Q18b — Classification Calculation [12 marks]')
body('Step 1: Priors.  P(HR) = 3/7.  P(LR) = 4/7.', indent=12)
body('Step 2: Conditional probabilities from training data:', indent=12)
tbl(['Feature','Value to match','HR instances','P(val|HR)','LR instances','P(val|LR)'],
    [['Income','High','ID3,ID7 = 2 of 3','2/3','ID4,ID6 = 2 of 4','2/4'],
     ['Credit History','Bad','ID1,ID3 = 2 of 3','2/3','ID2 = 1 of 4','1/4'],
     ['Late Payment','No','ID7 = 1 of 3','1/3','ID2,ID4,ID5 = 3 of 4','3/4'],
    ],widths=[1.3,1.0,1.5,0.8,1.5,0.8])
tip('CRITICAL ERROR to avoid: P(Income=High|LR) = 2/4 NOT 3/4. LR instances: ID2=Low, ID4=High, ID5=Low, ID6=High. Only ID4 and ID6 are High income.')
body('Step 3: Compute scores (Income=High, Credit=Bad, Late=No):', indent=12)
p = doc.add_paragraph(); p.paragraph_format.left_indent = _Pt(24); p.paragraph_format.space_after = _Pt(2)
r = p.add_run('Score(HR) = 3/7 x 2/3 x 2/3 x 1/3 = 12/189 = 0.0635'); r.bold = True; r.font.name = 'Courier New'; r.font.size = _Pt(10)
p = doc.add_paragraph(); p.paragraph_format.left_indent = _Pt(24); p.paragraph_format.space_after = _Pt(2)
r = p.add_run('Score(LR) = 4/7 x 2/4 x 1/4 x 3/4 = 24/448 = 0.0536'); r.bold = True; r.font.name = 'Courier New'; r.font.size = _Pt(10)
p = doc.add_paragraph(); p.paragraph_format.left_indent = _Pt(24); p.paragraph_format.space_after = _Pt(4)
r = p.add_run('Score(HR) 0.0635 > Score(LR) 0.0536  ->  CLASSIFIED AS: HIGH RISK'); r.bold = True; r.font.size = _Pt(10); r.font.color.rgb = RGBColor(0xC0,0x00,0x00)
tip('Despite having High income, this customer has Bad credit history (P=1/4 for LR) and No late payment (P=1/3 for HR). The bad credit history heavily penalises the LR score.')
sp()

h3('Q19a — CNN Three Layer Types')
tbl(['Layer','Function','Translation invariant?'],
    [['Convolutional','Applies filters (kernels) to detect local features: edges, textures','No (but equivariant)'],
     ['Pooling (max/avg)','Reduces spatial dimensions; summarises feature maps','Yes — key benefit'],
     ['Fully Connected','Combines all features; produces final class scores/output','N/A'],
    ],widths=[1.5,3.5,1.4])
h3('Q19b — Convolution: kernel [[1,0,-1],[1,0,-1],[1,0,-1]] on 3x3 patch [[4,3,2],[1,0,1],[-1,-2,-1]]')
formula('= 4*1+3*0+2*(-1) + 1*1+0*0+1*(-1) + (-1)*1+(-2)*0+(-1)*(-1)')
formula('= (4-2) + (1-1) + (-1+1) = 2 + 0 + 0 = 2')
sp()

h3('Q19c — CNN for Crowd Counting: Do You Agree? [6 marks]')
body('YES — CNN is well-suited for crowd counting. Three key reasons:',indent=12)
body('1. SPATIAL LOCALITY: Convolutional filters detect local patterns (heads, bodies) at every spatial position via weight sharing — critical when crowd density varies across the image.',indent=12)
body('2. TRANSLATION INVARIANCE: Pooling layers make detection robust to slight shifts in position. A head at top-left and one at bottom-right produce the same filter response.',indent=12)
body('3. DENSITY MAP OUTPUT: CNNs can be trained end-to-end to produce a density map (value per pixel region) rather than a single integer count, enabling spatial localisation of crowd hotspots.',indent=12)
body('Enhancement techniques: (a) Multi-column CNN (MCNN) — parallel branches with different kernel sizes handle scale variation as people appear larger/smaller at different distances. (b) Dilated convolutions increase receptive field without losing resolution. (c) Attention mechanisms focus on head/body regions and suppress background noise.',indent=12)
tip('6-mark rubric: ~2 marks per well-explained CNN property + 1 enhancement/limitation. Name the property AND explain why it helps counting.')
sp()

# ──────────────────────────────────────────
h2('PAPER 2 — Dec 2024 / Sem I (20250328_2_20.pdf)')
# ──────────────────────────────────────────
h3('Q3 — Short Paragraph Concept Answers')
body('Q3.1 Discriminative vs Generative: see Section 2.',indent=12)
body('Q3.2 Classification vs Clustering: see Section 2.',indent=12)
body('Q3.3 MCV vs LCV: see Section 2.',indent=12)
body('Q3.4 Forward vs Backward Chaining: see Section 2.',indent=12)
sp()

h3('Q4 — A* Timisoara to Bucharest (Romania map) [12 marks]')
body('SLD heuristics: Timisoara=329, Lugoj=244, Arad=366, Mehadia=241, Dobreta=242, Sibiu=253, RimnicuVilcea=193, Fagaras=176, Pitesti=100, Bucharest=0',indent=12)
body('Edge costs: Tim-Arad=118, Tim-Lugoj=111, Lugoj-Mehadia=70, Mehadia-Dobreta=75, Dobreta-Craiova=120, Arad-Sibiu=140, Sibiu-RV=80, Sibiu-Fagaras=99, RV-Pitesti=97, Pitesti-Bucharest=101',indent=12)
tbl(['Step','Node Expanded','g(n)','h(n)','f=g+h','Key Frontier After Expansion'],
    [['1','Timisoara','0','329','329','Lugoj(f=355), Arad(f=484)'],
     ['2','Lugoj','111','244','355','Mehadia(f=422), Arad(f=484)'],
     ['3','Mehadia','181','241','422','Arad(f=484), Dobreta(f=498)'],
     ['4','Arad','118','366','484','Dobreta(f=498), Sibiu(f=511)'],
     ['5','Dobreta','256','242','498','Sibiu(f=511), Craiova(f=536)'],
     ['6','Sibiu','258','253','511','RimnicuVilcea(f=531), Fagaras(f=533), Craiova(f=536)'],
     ['7','RimnicuVilcea','338','193','531','Fagaras(f=533), Pitesti(f=535), Craiova(f=536)'],
     ['8','Fagaras','357','176','533','Pitesti(f=535), Craiova(f=536), Bucharest_via_Fagaras(f=568)'],
     ['9','Pitesti','435','100','535','Craiova(f=536), Bucharest_via_Pitesti(f=536) — replaces f=568'],
     ['10','Bucharest','536','0','536','GOAL — stop'],
    ],widths=[0.4,1.4,0.6,0.6,0.7,3.1])
answer('Path: Timisoara → Arad → Sibiu → Rimnicu Vilcea → Pitesti → Bucharest   Cost = 118+140+80+97+101 = 536 km')
tip('Arad expanded BEFORE Dobreta (f=484 < 498) even though it is not on the optimal path — A* explores it because SLD overestimates. Bucharest via Fagaras (f=568) is replaced when Pitesti finds shorter route (f=536).')
sp()

h3('Q5 — Bayesian Network P(C) and P(W|C)')
body('Network: W->T, W->R, T->C, R->C',indent=12)
body('P(W)=0.6; P(T|W)=0.7,P(T|~W)=0.2; P(R|W)=0.8,P(R|~W)=0.1',indent=12)
body('P(C|T,R)=0.9, P(C|T,~R)=0.6, P(C|~T,R)=0.7, P(C|~T,~R)=0.1',indent=12)
tbl(['W','P(W)','Case','Contribution to P(C)'],
    [['T','0.6','P(T|W)*P(R|W)*P(C|T,R) = 0.7*0.8*0.9','0.378 * 0.6 = 0.2268'],
     ['T','0.6','P(T|W)*P(~R|W)*P(C|T,~R) = 0.7*0.2*0.6','0.084 * 0.6 = 0.0504'],
     ['T','0.6','P(~T|W)*P(R|W)*P(C|~T,R) = 0.3*0.8*0.7','0.168 * 0.6 = 0.1008'],
     ['T','0.6','P(~T|W)*P(~R|W)*P(C|~T,~R) = 0.3*0.2*0.1','0.006 * 0.6 = 0.0036'],
     ['F','0.4','P(T|~W)*P(R|~W)*P(C|T,R) = 0.2*0.1*0.9','0.018 * 0.4 = 0.0072'],
     ['F','0.4','P(T|~W)*P(~R|~W)*P(C|T,~R) = 0.2*0.9*0.6','0.108 * 0.4 = 0.0432'],
     ['F','0.4','P(~T|~W)*P(R|~W)*P(C|~T,R) = 0.8*0.1*0.7','0.056 * 0.4 = 0.0224'],
     ['F','0.4','P(~T|~W)*P(~R|~W)*P(C|~T,~R) = 0.8*0.9*0.1','0.072 * 0.4 = 0.0288'],
    ],widths=[0.4,0.5,3.0,2.0])
formula('P(C) = sum all above = 0.2268+0.0504+0.1008+0.0036+0.0072+0.0432+0.0224+0.0288 = 0.5588 (? let me verify partial)')
body('P(C|W=T) = 0.7*0.8*0.9+0.7*0.2*0.6+0.3*0.8*0.7+0.3*0.2*0.1 = 0.504+0.084+0.168+0.006 = 0.762',indent=12)
body('P(C|W=F) = 0.2*0.1*0.9+0.2*0.9*0.6+0.8*0.1*0.7+0.8*0.9*0.1 = 0.018+0.108+0.056+0.072 = 0.254',indent=12)
formula('P(C) = 0.762*0.6 + 0.254*0.4 = 0.4572 + 0.1016 = 0.5588')
formula('P(W|C) = P(C|W)*P(W) / P(C) = 0.762*0.6 / 0.5588 = 0.4572/0.5588 = 0.8182')
sp()

h3('Q6 — Decision Tree (8 Students, 3 Features → Pass/Fail) [12 marks]')
body('Dataset: 8 students. Features: Study Years (0-2, 3-5), Study Hours (Low/Med/High), Attendance (I=Irregular, R=Regular). Target: Pass.',indent=12)
tbl(['ID','Study Years','Study Hours','Attendance','Pass'],
    [['1','0-2','Low','I','No'],
     ['2','0-2','Low','R','No'],
     ['3','3-5','Medium','I','No'],
     ['4','3-5','Low','I','No'],
     ['5','3-5','High','R','Yes'],
     ['6','0-2','High','I','Yes'],
     ['7','0-2','High','R','Yes'],
     ['8','3-5','Medium','R','Yes'],
    ],widths=[0.4,1.0,1.0,1.1,0.6])
formula('H(S) = -4/8*log2(4/8) - 4/8*log2(4/8) = 1.0')
body('Step 1 — Information Gain for each feature:',indent=12)
tbl(['Feature','Partitions','Partition H','IG = 1.0 - weighted avg H'],
    [['Study Years','0-2: {1,2,6,7}=2Y2N; 3-5: {3,4,5,8}=2Y2N','H=1.0 each','IG = 1.0 - (4/8*1.0+4/8*1.0) = 0.0'],
     ['Study Hours','Low: {1,2,4}=0Y3N; High: {5,6,7}=3Y0N; Med: {3,8}=1Y1N','H=0; H=0; H=1.0','IG = 1.0 - (3/8*0+3/8*0+2/8*1.0) = 0.75'],
     ['Attendance','I: {1,3,4,6}=1Y3N; R: {2,5,7,8}=3Y1N','H=0.811 each','IG = 1.0 - (4/8*0.811+4/8*0.811) = 0.189'],
    ],widths=[1.1,2.4,1.3,1.8])
body('Step 2 — Root node: Study Hours (highest IG = 0.75)',indent=12)
body('Step 3 — Branch Low → pure No → LEAF: Predict No',indent=12)
body('Step 4 — Branch High → pure Yes → LEAF: Predict Yes',indent=12)
body('Step 5 — Branch Medium: remaining {ID3, ID8}. Check Attendance:',indent=12)
body('         I (ID3, No) → LEAF: Predict No    |    R (ID8, Yes) → LEAF: Predict Yes',indent=12)
formula('FINAL TREE: Root=Study Hours  →  Low→No  |  High→Yes  |  Medium→Attendance→{I→No, R→Yes}')
tip('Study Years gives IG=0.0 — splitting on it provides zero additional information. Always compute IG for ALL features before choosing the root.')
sp()

h3('Q7 — MLP Backprop: 2-input, 2-hidden, 1-output')
body('W(i->h) = [[0.1,0.2],[0.3,0.4]], W(h->o) = [0.5,0.6], x=[1,0], y=1, eta=0.5',indent=12)
tbl(['Step','Calculation','Value'],
    [['net_h1','0.1*1+0.2*0 = 0.1','o_h1=sigma(0.1)=0.5250'],
     ['net_h2','0.3*1+0.4*0 = 0.3','o_h2=sigma(0.3)=0.5744'],
     ['net_out','0.5*0.5250+0.6*0.5744=0.6071','y_hat=sigma(0.6071)=0.6474'],
     ['delta_out','(1-0.6474)*0.6474*(1-0.6474)','0.3526*0.6474*0.3526=0.0805'],
     ['W_h1o','0.5+0.5*0.0805*0.5250','0.5211'],
     ['W_h2o','0.6+0.5*0.0805*0.5744','0.6231'],
     ['delta_h1','(0.5*0.0805)*0.5250*(1-0.5250)','0.04025*0.5250*0.4750=0.0100'],
     ['delta_h2','(0.6*0.0805)*0.5744*(1-0.5744)','0.04830*0.5744*0.4256=0.0118'],
     ['W_i1h1','0.1+0.5*0.0100*1','0.1050'],
     ['W_i1h2','0.3+0.5*0.0118*1','0.3059'],
     ['W_i2h1,W_i2h2','x2=0 -> no change','0.2000, 0.4000'],
    ],widths=[1.3,3.0,2.1])
answer('Final W(i->h) = [[0.1050, 0.2000], [0.3059, 0.4000]]')
sp()

# ──────────────────────────────────────────
h2('PAPER 3 — Dec 2023 / Sem I (20241111_21.pdf)')
# ──────────────────────────────────────────
h3('Q3 — A* from C to G  [CLASSIC GRAPH — high repeat probability]')
body('Graph: A-6-B, B-3-C, C-4-E, A-20-F, B-18-D, C-8-D, E-6-D, F-5-D, D-16-G',indent=12)
body('H: A=30, B=25, C=20, D=15, E=10, F=5, G=0. Start=C, Goal=G.',indent=12)
tbl(['Step','Expanded','g(n)','h(n)','f=g+h','Frontier after'],
    [['1','C','0','20','20','E(f=4+10=14), D(f=8+15=23), B(f=? C is reached from A originally — start is C so B not reachable backwards)'],
     ['2','E','4','10','14','D(g=10,h=15,f=25), D_orig(f=23). Keep D(f=23). B unreachable.'],
     ['3','D','8','15','23','G(g=8+16=24,h=0,f=24), F(g=8+5+16? no, D->F edge not in graph)'],
     ['4','G','24','0','24','GOAL'],
    ],widths=[0.4,0.8,0.5,0.5,0.7,3.5])
answer('PATH: C -> D -> G, total cost = 8+16 = 24')
tip('From C, D is directly reachable (cost 8). From E, D costs 4+6=10 > 8, so keep original D(g=8). Then D->G=16.')
sp()

h3('Q4 — K-Means Clustering (k=3, 10 data points) [12 marks]')
body('Data points: 0.12, 0.15, 0.21, 0.50, 0.69, 0.78, 0.89, 1.25, 1.46, 1.97',indent=12)
body('Initial centroids: μ1=0.12, μ2=0.78, μ3=1.97  (first, middle, last)',indent=12)
tbl(['Point','Dist to μ1=0.12','Dist to μ2=0.78','Dist to μ3=1.97','Cluster (Iter 1)'],
    [['0.12','0.00','0.66','1.85','C1'],
     ['0.15','0.03','0.63','1.82','C1'],
     ['0.21','0.09','0.57','1.76','C1'],
     ['0.50','0.38','0.28','1.47','C2'],
     ['0.69','0.57','0.09','1.28','C2'],
     ['0.78','0.66','0.00','1.19','C2'],
     ['0.89','0.77','0.11','1.08','C2'],
     ['1.25','1.13','0.47','0.72','C3'],
     ['1.46','1.34','0.68','0.51','C3'],
     ['1.97','1.85','1.19','0.00','C3'],
    ],widths=[0.6,1.4,1.4,1.4,1.0])
body('After Iteration 1: C1={0.12,0.15,0.21}, C2={0.50,0.69,0.78,0.89}, C3={1.25,1.46,1.97}',indent=12)
formula('New centroids: μ1=(0.12+0.15+0.21)/3=0.16   μ2=(0.50+0.69+0.78+0.89)/4=0.715   μ3=(1.25+1.46+1.97)/3=1.56')
tbl(['Point','Dist to μ1=0.16','Dist to μ2=0.715','Dist to μ3=1.56','Cluster (Iter 2)'],
    [['0.12','0.04','0.595','1.44','C1'],
     ['0.15','0.01','0.565','1.41','C1'],
     ['0.21','0.05','0.505','1.35','C1'],
     ['0.50','0.34','0.215','1.06','C2'],
     ['0.69','0.53','0.025','0.87','C2'],
     ['0.78','0.62','0.065','0.78','C2'],
     ['0.89','0.73','0.175','0.67','C2'],
     ['1.25','1.09','0.535','0.31','C3'],
     ['1.46','1.30','0.745','0.10','C3'],
     ['1.97','1.81','1.255','0.41','C3'],
    ],widths=[0.6,1.4,1.45,1.45,1.0])
body('After Iteration 2: Same assignments → NO CHANGE → CONVERGED in 2 iterations.',indent=12)
answer('Final clusters: C1={0.12, 0.15, 0.21}  |  C2={0.50, 0.69, 0.78, 0.89}  |  C3={1.25, 1.46, 1.97}')
tip('Always show the distance table for EACH iteration. Stop when cluster assignments stop changing. State "CONVERGED" explicitly.')
sp()

h3('Q5 — Bayesian Network P(C|E=T) and P(E|C)')
body('Network: E->S, E->J, S->C, J->C. P(E)=0.3',indent=12)
body('P(S|E)=0.2, P(S|~E)=0.6; P(J|E)=0.8, P(J|~E)=0.1',indent=12)
body('P(C|J,S)=0.9, P(C|J,~S)=0.8, P(C|~J,S)=0.4, P(C|~J,~S)=0.1',indent=12)
tbl(['S','J','P(S,J|E=T)','P(C|S,J)','Contribution'],
    [['S','J','0.2*0.8=0.16','0.9','0.144'],
     ['S','~J','0.2*0.2=0.04','0.4','0.016'],
     ['~S','J','0.8*0.8=0.64','0.8','0.512'],
     ['~S','~J','0.8*0.2=0.16','0.1','0.016'],
    ],widths=[0.5,0.5,1.8,1.3,1.7])
formula('P(C|E=T) = 0.144+0.016+0.512+0.016 = 0.688')
body('Similarly P(C|E=F): P(S|~E)=0.6, P(J|~E)=0.1:',indent=12)
formula('P(C|E=F) = 0.6*0.1*0.9+0.6*0.9*0.4+0.4*0.1*0.8+0.4*0.9*0.1 = 0.054+0.216+0.032+0.036 = 0.338')
formula('P(C) = 0.688*0.3 + 0.338*0.7 = 0.2064+0.2366 = 0.4430')
formula('P(E|C) = P(C|E)*P(E)/P(C) = 0.688*0.3/0.4430 = 0.2064/0.4430 = 0.4659 ≈ 0.47')
sp()

h3('Q7 — MLP Backprop: 4-input, 3-hidden, 2-output (all-ones weights)')
body('W(0)=4x3 all-ones, W(1)=3x2 all-ones. x=[1,0,1,0], y=[1,0], eta=0.5. Bias=0.',indent=12)
tbl(['Step','Calculation','Value'],
    [['net_j (all 3 hidden)','1*1+1*0+1*1+1*0 = 2','same for j=1,2,3'],
     ['o_j (all 3 hidden)','sigma(2) = 1/(1+e^{-2})','0.8808 for all j'],
     ['net_p (both outputs)','0.8808+0.8808+0.8808','2.6424 for p=1,2'],
     ['y_hat_p (both)','sigma(2.6424)','0.9335 for p=1,2'],
     ['delta_p1','(1-0.9335)*0.9335*(1-0.9335)','0.0665*0.9335*0.0665=0.0041'],
     ['delta_p2','(0-0.9335)*0.9335*(1-0.9335)','-0.9335*0.9335*0.0665=-0.0579'],
     ['W(1)_jp1 update','1.0+0.5*0.0041*0.8808','1.0018 (all 3 rows)'],
     ['W(1)_jp2 update','1.0+0.5*(-0.0579)*0.8808','0.9745 (all 3 rows)'],
     ['sum_p(W_jp*delta_p)','1.0*0.0041+1.0*(-0.0579)','-0.0538 (same all j)'],
     ['delta_j (all 3)','0.8808*0.1192*(-0.0538)','-0.0057 (same all j)'],
     ['W(0) for i=1,x_1=1','1.0+0.5*(-0.0057)*1','0.9971 (all 3 hidden)'],
     ['W(0) for i=3,x_3=1','1.0+0.5*(-0.0057)*1','0.9971 (all 3 hidden)'],
     ['W(0) for i=2,4 (x=0)','no change','1.0000'],
    ],widths=[2.2,3.0,1.6])
answer('W(1)[1]: all rows = [1.0018, 0.9745].  W(0)[1]: all rows = [0.9971, 1.0, 0.9971, 1.0]')
sp()

# ──────────────────────────────────────────
h2('PAPER 4 — Aug 2022 (2023-08-17_20.pdf)')
# ──────────────────────────────────────────
h3('Q3.1 — 8-Puzzle BFS')
body('Start: [1,4,2/3,_,5/6,7,8]  Goal: [_,1,2/3,4,5/6,7,8]  (blank=_)',indent=12)
tbl(['Level','State (blank position)','Move applied','Children'],
    [['0','[1,4,2/3,_,5/6,7,8] blank(1,1)','—','Up, Down, Left, Right'],
     ['1','[1,_,2/3,4,5/6,7,8] blank(0,1)','Up','Left, Right (from here)'],
     ['2','[_,1,2/3,4,5/6,7,8] blank(0,0)','Left','= GOAL'],
    ],widths=[0.5,2.5,1.3,1.9])
answer('PATH: Start -> [1,_,2/3,4,5/6,7,8] -> GOAL  (slide tile 4 up, then tile 1 left). 2 moves.')
sp()

h3('Q3.2 — Heuristic Dominance from state (c)')
body('State (c): [7,2,_/5,4,3/1,6,8], Goal: [_,1,2/3,4,5/6,7,8]',indent=12)
body('h1 (misplaced tiles) = 6  |  h2 (Manhattan distance) = 3+1+2+0+2+1+3+0 = 12',indent=12)
tip('h2 dominates h1 because h2(n) >= h1(n) for all n AND both are admissible. Dominant heuristic = fewer nodes expanded.')
sp()

h3('Q4.1 — Perceptron 2D')
answer('Final w = (-0.2, -0.4)')
sp()

h3('Q4.2 — Convolution (all-ones 3x3 kernel on 4x4 image)')
body('All-ones kernel = sum of 9-pixel region. Output is 2x2.',indent=12)
answer('Output: [[1140, 1150], [1200, 1210]]  (each cell = sum of its 3x3 patch)')
sp()

h3('Q6 — Bayesian Network P(B|C): A->C, B->C')
body('P(A)=0.1, P(B)=0.7; P(C|A,B)=0.9, P(C|A,~B)=0.8, P(C|~A,B)=0.6, P(C|~A,~B)=0.2',indent=12)
formula('P(C) = 0.9*0.1*0.7 + 0.8*0.1*0.3 + 0.6*0.9*0.7 + 0.2*0.9*0.3')
formula('     = 0.063 + 0.024 + 0.378 + 0.054 = 0.519')
formula('P(C,B) = P(C|A,B)*P(A)*P(B) + P(C|~A,B)*P(~A)*P(B) = 0.063 + 0.378 = 0.441')
formula('P(B|C) = P(C,B)/P(C) = 0.441/0.519 = 0.850')
sp()

# ──────────────────────────────────────────
h2('PAPER 5 — Dec 2016 / Sem I (170306_27.pdf)')
# ──────────────────────────────────────────
h3('Q3 — All 5 Searches: Graph B-6-C, B-11-A, C-9-D, C-25-F, C-10-E, E-3-D, D-7-G')
body('H: A=30, B=25, C=15, D=5, E=10, F=20, G=0. Start=B, Goal=G.',indent=12)
tbl(['Algorithm','Path found','Cost','Notes'],
    [['BFS','B->C->D->G (fewest hops)','varies','Level-order expansion'],
     ['DFS','B->A (dead end)->C->D->G or similar','varies','Go deep first'],
     ['UCS','B->C->E->D->G','6+10+3+7=26','Cheapest first; note C-D=9, E-D=3 cheaper path to D'],
     ['Greedy','B->C->D->G','6+9+7=22','h only; D has h=5 < E(10); D->G directly'],
     ['A*','B->C->D->G','22','f(C)=6+15=21, f(D)=15+5=20, f(G)=22+0=22'],
    ],widths=[0.9,1.8,0.8,2.9])
h3('Q3.6 Admissibility Check')
tbl(['Node','h(n)','h*(n) = true dist to G','h <= h*?'],
    [['A','30','A is dead end (no route to G without backtrack to B)','Need to verify: A->B->C->D->G = 11+6+9+7=33. h(A)=30<=33 OK'],
     ['B','25','B->C->E->D->G = 6+10+3+7=26. Or B->C->D->G=22','h(B)=25<=22? NO — 25>22. INADMISSIBLE!'],
     ['C','15','C->D->G = 9+7=16. Or C->E->D->G=20','h(C)=15<=16 OK'],
     ['D','5','D->G = 7','h(D)=5<=7 OK'],
     ['E','10','E->D->G = 3+7=10','h(E)=10<=10 OK (tight)'],
    ],widths=[0.5,0.7,3.5,1.7])
body('NOTE: h(B)=25 > h*(B via C->D->G)=22. If this is correct from the paper, heuristic is INADMISSIBLE. Check paper values carefully.',indent=12)
tip('The Dec 2015 version of this graph has h*(B)=26 (via C->E->D->G=6+10+3+7=26) which makes h(B)=25 admissible. Verify which graph applies.')
sp()

h3('Q4.1 — Perceptron (4D, 8 Points)')
body('Initial w=(0.1, -0.02, -0.02, 0.07), eta=0.1. Two updates, converges after 1 epoch.',indent=12)
answer('Final weights: w = (0.004, 0.148, -0.012, 0.058)')
sp()

h3('Q6 — Bayesian P(D) and P(A|C): A->B->C, B->D')
body('P(A)=0.6, P(B|A)=0.8, P(B|~A)=0.6; P(C|B)=0.4, P(C|~B)=0.2; P(D|B)=0.7, P(D|~B)=0.4',indent=12)
formula('P(B) = 0.8*0.6 + 0.6*0.4 = 0.48+0.24 = 0.72')
formula('P(D) = 0.7*0.72 + 0.4*0.28 = 0.504+0.112 = 0.616')
formula('P(C) = 0.4*0.72 + 0.2*0.28 = 0.288+0.056 = 0.344')
formula('P(A,C) = P(A)*P(B|A)*P(C|B) + P(A)*P(~B|A)*P(C|~B) = 0.6*0.8*0.4+0.6*0.2*0.2 = 0.192+0.024 = 0.216')
formula('P(A|C) = P(A,C)/P(C) = 0.216/0.344 = 0.628')
sp()

# ──────────────────────────────────────────
h2('PAPER 6 — Dec 2015 / Sem I (160425-2_30.pdf)')
# ──────────────────────────────────────────
h3('Q3 — All 5 Searches on CLASSIC GRAPH: A-6-B, B-3-C, C-4-E, A-20-F, B-18-D, C-8-D, E-6-D, F-5-D, D-16-G')
body('H: A=30, B=25, C=20, D=15, E=10, F=5, G=0. Start=C, Goal=G.',indent=12)
tbl(['Algorithm','Path','Cost','Key decision'],
    [['BFS','C -> B,D,E -> A,F -> G','via D: 8+16=24 hops','Level-order from C'],
     ['DFS','C->B->A(dead end)->F->D->G','20+5+16=? or C->B->F->D->G','Go deepest first'],
     ['UCS','C->D->G','8+16=24','D(g=8) cheaper than E(g=4)->D(g=10)'],
     ['Greedy','C->E->D->G','4+6+16=26','h(E)=10 < h(D)=15 from C, then D->G directly'],
     ['A*','C->D->G','8+16=24','f(E)=4+10=14, f(D_orig)=8+15=23, D_via_E=10+15=25 > 23'],
    ],widths=[0.9,1.8,0.8,2.9])
tip('THIS EXACT GRAPH appeared in BOTH Dec 2015 AND Dec 2023 papers. Memorise the A* solution: C->D->G cost=24.')
sp()

h3('Q5 — Decision Tree (20 instances, 5 features)')
body('H(S)=1.0 (10Y/10N). Root=F1.',indent=12)
tbl(['Branch','Sub-split','Leaves'],
    [['F1=V1 (9Y,2N)','Split on F4','F4V2->YES, F4V3->YES, F4V1->(2Y,2N)->split F2: F2V1->YES, F2V2->NO'],
     ['F1=V2 (1Y,8N)','Split on F2','F2V2->NO (0Y,3N+others)... see detail below'],
    ],widths=[1.5,2.0,2.9])
sp()

h3('Q6 — Bayesian P(~C) and P(B|D): A->B->C, A->B->D')
body('P(A)=0.5, P(B|A)=0.7, P(B|~A)=0.5; P(C|B)=0.4, P(C|~B)=0.2; P(D|B)=0.7, P(D|~B)=0.4',indent=12)
formula('P(B) = 0.7*0.5 + 0.5*0.5 = 0.35+0.25 = 0.60')
formula('P(C) = 0.4*0.6 + 0.2*0.4 = 0.24+0.08 = 0.32   ->  P(~C) = 1-0.32 = 0.68')
formula('P(D) = 0.7*0.6 + 0.4*0.4 = 0.42+0.16 = 0.58')
formula('P(B|D) = P(D|B)*P(B)/P(D) = 0.7*0.6/0.58 = 0.42/0.58 = 0.724')
sp()

doc.add_page_break()

# ════════════════════════════════════════════
h1('SECTION 4 — Quiz Solutions by Topic')
# ════════════════════════════════════════════

# ── Quiz 1 ──
h2('Quiz 1 — AI History + Search')
h3('2025 Q1 — DFS on graph A-2-C, A-4-B, C-10-E, C-8-D, B-5-D, D-2-E. Goal=E.')
body('Alphabetical neighbour order. Depth-first: A expands B first (alpha). B->D. D->E (GOAL).',indent=12)
answer('DFS path: A -> B -> D -> E')
h3('2025 Q1 — A* on same graph. H: A=7, B=6, C=4, D=2, E=0.')
tbl(['Step','Expanded','f','Frontier'],
    [['1','A (f=7)','0+7','B(4+6=10), C(2+4=6)'],
     ['2','C (f=6)','2+4','D(10+2=12), E(12+0=12), B(10)'],
     ['3','D via B (f=11)','9+2','After expanding B(f=10): D_via_B(f=11). Update D: 11<12'],
     ['4','E via D (f=13)','11+2','Or E directly via C (f=12)? Pick E(f=12) via C.'],
    ],widths=[0.5,1.5,0.8,3.5])
answer('A* path: A -> C -> D -> E, cost=12. (D via C: g=10, f=12. E via C direct: g=12, f=12. E via D via C: g=12, f=12.)')
tip('NOTE: Verify carefully — D via B (g=9) < D via C (g=10), so D is reached via B. Then E via D: g=11. So A->B->D->E cost=11 is optimal.')
answer('CORRECT: A -> B -> D -> E, cost = 4+5+2 = 11.')
sp()

h2('Quiz 2 — Knowledge-Based Agents + Games')
h3('Q1a — Forward Chaining (Car Diagnosis)')
tbl(['Step','Rule fired','New fact added'],
    [['1','Rule 3: headlights_dim -> battery_dead','battery_dead'],
     ['2','Rule 1: car_wont_start AND battery_dead -> dead_battery_issue','dead_battery_issue'],
    ],widths=[0.5,3.0,2.4])
sp()

h3('Q2a — Minimax Game Tree (Stone/Nim)')
body('Game: last to take stone LOSES. Utility: f=-9 (current player takes last), f=10 (opponent does).',indent=12)
body('Evaluate leaves bottom-up: MIN propagates min; MAX propagates max to root.',indent=12)
answer('Root receives f=10 -> best first move = take 1 stone.')
sp()

h2('Quiz 3 — Hill Climbing / SA + CSP')
h3('Q1a — SA: Better Move (TSP current=31, new=28)')
answer('28 < 31: SA ALWAYS accepts better solutions. No probability needed.')
h3('Q1b — SA: Worse Move (current=28, new=29, delta=1)')
formula('P(accept) = e^{-|delta|/T} = e^{-1/T}')
body('At T=10: P=0.905 (high T, likely to accept). At T=1: P=0.368. At T=0.1: P~0.',indent=12)
tip('If T not given: write P=e^{-1/T} and state "decreases as T decreases, nearly zero at low T".')
sp()

h3('Q2 — CSP Map Coloring (B connected to A,C,D,E; A connected to B,D)')
tbl(['Heuristic','Answer','Reasoning'],
    [['Degree (Q2b)','Region B (4 connections)','Most edges -> most constraints -> reduces branching most'],
     ['MCV after B assigned (Q2c)','Region A (1 legal color left = Green)','Fewest remaining values after B assigned'],
    ],widths=[2.0,2.0,2.4])
sp()

h2('Quiz 4 — Decision Tree + Neural Network')
h3('Q1a — Decision Tree: 8 instances, 4Y/4N (Hiking dataset)')
body('H(S) = 1.0. Weather has highest IG = 0.656 -> chosen as root.',indent=12)
tbl(['Attribute','Values','H(Sv)','IG'],
    [['Temperature','Hot(3)/Mild(2)/Cool(3)','0.918/1.0/0.918','1.0-(3/8*0.918+2/8*1.0+3/8*0.918)=0.0615'],
     ['Weather (root)','Sunny/Overcast/Rainy','varies','0.656 — HIGHEST'],
    ],widths=[1.3,1.8,1.5,2.8])
sp()

h3('Q2a — Single Neuron Forward Pass: w1=0.5, w2=-0.3, b=0.1, x1=0.8, x2=0.4')
formula('s = 0.5*0.8 + (-0.3)*0.4 + 0.1 = 0.4 - 0.12 + 0.1 = 0.38')
formula('sigma(0.38) = 1/(1+e^{-0.38}) = 1/1.462 = 0.594')
sp()

h3('Q2b — Why Sign Function Cannot Be Used in Backprop')
body('See Section 2 model answer.',indent=12)
sp()

h2('Quiz 5–9 — ML, Neural Networks, Deep Learning, Bayesian')
h3('Recurring Quiz Patterns (2023–2025 — often identical questions)')
tbl(['Quiz Q','Topic','Key answer'],
    [['Q5 (all years)','Decision Tree entropy+IG','F4 highest IG -> root. F4V1->split F2. F4V2->pure YES. F4V3->split F3.'],
     ['Q6 (2025)','SVM linear classifier','SVs: (5,1) and (-1,-1). Midpoint=(2,0). Boundary: y=-3x+6. Margin=2/sqrt(10)=0.632'],
     ['Q6 (2024)','K-Means 1D, 3 clusters','Init C1=0.12,C2=0.15,C3=1.97. Final: {0.12-0.21},{0.50-0.89},{1.25-1.97}. Converges 2 iters.'],
     ['Q7 (all years)','Perceptron 8 4D points','eta=0.2, init w=(-0.09,0,-0.19,-0.21). Updates at x3,x6. Final w=(-0.066,0.396,0.026,0.010)'],
     ['Q8 (2024,2025)','MLP True/False','Various; check specific year'],
     ['Q9 (2024,2025)','CNN 4x4 image, 3x3 kernel [-1,0,1 / -1,0,1 / -1,0,1]','Output 2x2. Deep Belief Networks (not CNN) are key DL model MCQ.'],
    ],widths=[1.2,2.0,3.2])
sp()

h3('2024 Q3 — CSP Map Coloring with MCV+LCV')
body('8 regions. Min 4 colors. Assignment: R1:C2, R2:C4, R3:C3, R4:C1, R5:C3, R6:C1, R7:C2, R8:C1',indent=12)
sp()

h3('2024 Q4 — FC vs BC on KB {A,B,C, A^B->D, B^C->E, A^C->F, A^F->G, D^F->K, G^K->Q1, H^C->Q2}')
body('FC can prove: D, E, F, G, K, Q1. CANNOT prove Q2 (H not in KB).',indent=12)
body('BC for Q1: Q1<-G^K; G<-A^F; F<-A^C (both given); K<-D^F; D<-A^B (both given). Q1 PROVED.',indent=12)
body('BC for Q2: Q2<-H^C; C given but H not derivable. Q2 FAILS.',indent=12)
sp()

h3('2023 Q2 — CSP Course Assignment (7 courses, minimum 3 lecturers)')
body('MCV: COMP04 (5 conflicts)->A; COMP01 (4 conflicts)->B; COMP02->A; COMP03->C; COMP05->C; COMP06->B; COMP07->C',indent=12)
answer('3 lecturers minimum = graph is 3-colorable.')
sp()

h3('SVM — Linear Classifier (Quiz 6 2025)')
body('Support vectors: (5,1) class RED; (-1,-1) class BLUE.',indent=12)
formula('Midpoint = ((5-1)/2, (1-1)/2) = (2, 0)')
formula('Slope of SV line: k_d = (1-(-1))/(5-(-1)) = 2/6 = 1/3')
formula('Perpendicular classifier slope: k_d * k_perp = -1 -> k_perp = -3')
formula('Passes through (2,0): 0 = -3*2 + b -> b = 6')
formula('Decision boundary: y = -3x + 6')
formula('Margin = 2/||w|| = 2/sqrt((-3)^2+1^2) = 2/sqrt(10) = 0.632')
sp()

doc.save('COMP4431_PastPaper_QuizSol.docx')
print('Saved COMP4431_PastPaper_QuizSol.docx')

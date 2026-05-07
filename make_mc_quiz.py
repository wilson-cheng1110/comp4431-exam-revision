from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── page margins ──────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)

# ── helpers ───────────────────────────────────────────────
def style(para, size=11, bold=False, italic=False, color=None):
    for run in para.runs:
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
    return para

def h1(t):
    p = doc.add_heading(t, 1)
    p.runs[0].font.color.rgb = RGBColor(0x1D, 0x40, 0xAF)
    return p

def h2(t):
    p = doc.add_heading(t, 2)
    p.runs[0].font.color.rgb = RGBColor(0x10, 0x59, 0x59)
    return p

def sp():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

def question(num, stem, options, correct, explanation):
    """num=int, stem=str, options=list[str] (4 items), correct='A'/'B'/'C'/'D', explanation=str"""
    QUESTIONS.append((num, stem, options, correct, explanation))

# ── question bank ─────────────────────────────────────────
QUESTIONS = []

# ══════════════════════════════════════════════════════════
# L01: Introduction to AI  (Q1-Q10)
# ══════════════════════════════════════════════════════════
question(1,
    "What is the key difference between a DISCRIMINATIVE and a GENERATIVE AI model?",
    ["A) Discriminative models require more training data than generative models",
     "B) Discriminative models learn P(Y|X); generative models learn the joint P(X,Y)",
     "C) Generative models are always supervised; discriminative models are always unsupervised",
     "D) Discriminative models can generate new images; generative models cannot"],
    "B",
    "Discriminative = P(Y|X) = classify. Generative = P(X,Y) = model the data distribution and can create new samples (e.g. GANs).")

question(2,
    "A thermostat that reads temperature and turns the heater on/off with no memory is best described as:",
    ["A) A goal-based agent",
     "B) A utility-based agent",
     "C) A simple reflex agent",
     "D) A model-based reflex agent"],
    "C",
    "Simple reflex agents react directly to the current percept via condition-action rules. No memory, no model of the world.")

question(3,
    "A chess game where both players can see the full board at all times is:",
    ["A) Partially observable and stochastic",
     "B) Fully observable and deterministic",
     "C) Fully observable and stochastic",
     "D) Partially observable and deterministic"],
    "B",
    "Both players see the complete board (fully observable). Given the same position, the same move always produces the same next state (deterministic).")

question(4,
    "Which of the following is NOT one of the 7 Dartmouth research directions from 1956?",
    ["A) Neural networks",
     "B) Natural language processing",
     "C) Deep learning and Transformer architectures",
     "D) Creativity and randomness"],
    "C",
    "Deep learning and Transformers are modern (2010s onwards). The 1956 Dartmouth workshop predated them by decades.")

question(5,
    "PAC learning states that the sample complexity satisfies n ≥ (1/ε)·ln(|H|/δ). Here, (1−δ) represents:",
    ["A) The maximum allowable error rate",
     "B) The probability with which the learned hypothesis has error ≤ ε",
     "C) The size of the hypothesis space",
     "D) The fraction of training examples that may be mislabelled"],
    "B",
    "PAC: with probability ≥ (1−δ), the error of h is ≤ ε. δ is the failure probability; (1−δ) is the confidence.")

question(6,
    "Which agent type selects actions by maximising an expected utility function, handling trade-offs between competing goals?",
    ["A) Simple reflex agent",
     "B) Model-based reflex agent",
     "C) Goal-based agent",
     "D) Utility-based agent"],
    "D",
    "Utility-based agents assign a real-valued score to states and choose actions that maximise expected utility.")

question(7,
    "A self-driving car operates in an environment where other drivers' actions are unpredictable. This environment is best classified as:",
    ["A) Fully observable and deterministic",
     "B) Stochastic (non-deterministic)",
     "C) Episodic and discrete",
     "D) Fully observable and episodic"],
    "B",
    "Stochastic: the outcome of an action is not fully determined by the action and the current state (other drivers introduce randomness).")

question(8,
    "Which learning paradigm requires labelled (input, output) pairs to learn a mapping from inputs to outputs?",
    ["A) Unsupervised learning",
     "B) Reinforcement learning",
     "C) Supervised learning",
     "D) Self-supervised learning"],
    "C",
    "Supervised learning: trained on labelled examples. Unsupervised: no labels. RL: learns from rewards, not direct labels.")

question(9,
    "A 'model-based reflex agent' differs from a 'simple reflex agent' by:",
    ["A) Choosing actions to maximise a utility function",
     "B) Maintaining an internal model of the world to handle partial observability",
     "C) Learning from past experience",
     "D) Using a planning algorithm to find optimal sequences"],
    "B",
    "Model-based reflex agents keep an internal state (world model) updated over time, allowing them to handle partially observable environments.")

question(10,
    "Which statement about a 'sequential' vs 'episodic' environment is correct?",
    ["A) Episodic: current action may affect future rewards. Sequential: each episode is independent",
     "B) Sequential: current action may affect all future decisions. Episodic: each episode is independent",
     "C) They are the same concept described with different terminology",
     "D) Sequential environments always have a finite number of steps"],
    "B",
    "Sequential (e.g. chess): decisions have long-term consequences. Episodic (e.g. spam filter): each classification is independent.")

# ══════════════════════════════════════════════════════════
# L02: Heuristic Search  (Q11-Q25)
# ══════════════════════════════════════════════════════════
question(11,
    "Which search algorithm guarantees finding the OPTIMAL path when all edge costs are EQUAL?",
    ["A) DFS",
     "B) BFS",
     "C) Greedy",
     "D) A* with h=0 only"],
    "B",
    "BFS finds fewest hops. When all costs equal, fewest hops = least total cost. For non-uniform costs, use UCS or A*.")

question(12,
    "UCS (Uniform Cost Search) expands nodes in what order?",
    ["A) Fewest hops from the start node",
     "B) Lowest h(n) — heuristic estimate to goal",
     "C) Lowest g(n) — cumulative path cost from start",
     "D) Lowest f(n) = g(n) + h(n)"],
    "C",
    "UCS = priority queue ordered by g(n). It is equivalent to Dijkstra's shortest-path algorithm.")

question(13,
    "Why is Greedy search NOT optimal?",
    ["A) It expands too many nodes",
     "B) It uses only h(n) and ignores the actual path cost g(n) already incurred",
     "C) Its heuristic must be inadmissible",
     "D) It cannot find any solution in weighted graphs"],
    "B",
    "Greedy only looks ahead (heuristic), not back (cost so far). It may race toward the goal via an expensive path.")

question(14,
    "A heuristic h(n) is ADMISSIBLE if and only if:",
    ["A) h(n) ≥ h*(n) for all n",
     "B) h(n) = h*(n) for all n",
     "C) h(n) ≤ h*(n) for all n",
     "D) h(n) = 0 for all n"],
    "C",
    "Admissible = never overestimates the true cost h*(n). This guarantees A* finds the optimal path.")

question(15,
    "In A*, if a cheaper path to an already-explored (closed) node is found, you should:",
    ["A) Ignore it — the node is already expanded, never revisit",
     "B) Restart the entire search from the start node",
     "C) Update the node's g value and re-add it to the open list",
     "D) Discard the new path only if the heuristic is consistent"],
    "C",
    "A* must update and re-queue if a cheaper path is found — otherwise it may not find the optimal solution.")

question(16,
    "BFS uses a _______ ; DFS uses a _______.",
    ["A) Stack; Queue",
     "B) Priority queue; Stack",
     "C) Queue; Stack",
     "D) Queue; Priority queue"],
    "C",
    "BFS = FIFO queue (explores level by level). DFS = LIFO stack (dives deep first).")

question(17,
    "Which algorithm is COMPLETE but NOT OPTIMAL?",
    ["A) A* with admissible heuristic",
     "B) UCS on a graph with non-negative costs",
     "C) BFS on a finite graph with non-uniform costs",
     "D) DFS on a finite graph without cycles"],
    "C",
    "BFS is complete (finds a solution if one exists) but not optimal for non-uniform costs — it finds fewest hops, not least cost.")

question(18,
    "When h(n) = 0 for all nodes, A* behaves identically to:",
    ["A) Greedy search",
     "B) DFS",
     "C) UCS",
     "D) BFS"],
    "C",
    "f(n) = g(n) + 0 = g(n). A* reduces to UCS — expands purely by path cost.")

question(19,
    "h2 DOMINATES h1 means:",
    ["A) h2(n) ≤ h1(n) for all n",
     "B) h2(n) ≥ h1(n) for all n, and both h1 and h2 are admissible",
     "C) h2 always finds the solution faster regardless of admissibility",
     "D) h2 requires less memory than h1"],
    "B",
    "Dominant heuristic: more informed (h2 ≥ h1 everywhere) while both remain admissible. A* with h2 expands fewer nodes.")

question(20,
    "Manhattan distance dominates misplaced-tile count as an 8-puzzle heuristic because:",
    ["A) It is faster to compute",
     "B) It is inadmissible, providing a tighter upper bound",
     "C) It is always ≥ misplaced-tile count and both are admissible",
     "D) It accounts for diagonal moves in the puzzle"],
    "C",
    "Each misplaced tile contributes at least 1 to Manhattan distance. So Manhattan ≥ misplaced. Both never overestimate true cost.")

question(21,
    "In the Romania A* problem, Arad is expanded even though it is not on the optimal path. Why?",
    ["A) The heuristic is inadmissible at Arad",
     "B) Arad's f value is the lowest among all frontier nodes at that expansion step",
     "C) Arad is adjacent to the goal node Bucharest",
     "D) A* always expands all nodes adjacent to the start"],
    "B",
    "A* expands the node with lowest f. At step 4, Arad(f=484) < Dobreta(f=498) — Arad must be expanded first even though it is off-path.")

question(22,
    "A consistent (monotone) heuristic satisfies the triangle inequality:",
    ["A) h(n) ≤ h*(n) for all n",
     "B) h(n) ≥ cost(n, n') + h(n') for all successors n'",
     "C) h(n) ≤ cost(n, n') + h(n') for all successors n'",
     "D) h(n) = cost(n, n') + h(n') for the optimal successor"],
    "C",
    "Consistency: h(n) ≤ edge_cost(n→n') + h(n'). Consistent → admissible (not vice versa). Prevents re-expansion of nodes.")

question(23,
    "DFS without visited-state tracking on a graph (not a tree) can:",
    ["A) Find the optimal solution faster than A*",
     "B) Loop infinitely in cycles, never terminating",
     "C) Always expand fewer nodes than BFS",
     "D) Guarantee finding a path if the branching factor is finite"],
    "B",
    "DFS without cycle detection will follow A→B→A→B... forever on a graph with cycles. Always track visited states on graphs.")

question(24,
    "If all edge costs are equal to 1, which pair of searches always finds exactly the same path?",
    ["A) BFS and DFS",
     "B) BFS and UCS",
     "C) Greedy and A*",
     "D) UCS and DFS"],
    "B",
    "Uniform cost 1: fewest hops = cheapest path. Both BFS (min hops) and UCS (min cost) agree.")

question(25,
    "The open list in A* contains ___ nodes; the closed list contains ___ nodes.",
    ["A) Expanded nodes; unexpanded nodes",
     "B) Generated but not yet expanded nodes; already expanded nodes",
     "C) All visited nodes; all goal candidates",
     "D) Nodes with f > threshold; nodes with f ≤ threshold"],
    "B",
    "Open = frontier (generated, waiting to be expanded). Closed = already processed (don't re-expand unless cheaper path found).")

# ══════════════════════════════════════════════════════════
# L03: Knowledge-Based Agents  (Q26-Q35)
# ══════════════════════════════════════════════════════════
question(26,
    "Which statement about TELL and ASK on a Knowledge Base is correct?",
    ["A) TELL queries the KB; ASK adds new facts to the KB",
     "B) TELL adds sentences; ASK checks whether a sentence is entailed by the KB",
     "C) Both TELL and ASK add new facts — ASK adds the answer automatically",
     "D) ASK removes sentences from the KB that are no longer relevant"],
    "B",
    "TELL = add a sentence. ASK = query whether something is entailed. ASK does NOT modify the KB.")

question(27,
    "Forward chaining is best described as:",
    ["A) Goal-driven — works backward from the goal to find supporting facts",
     "B) Data-driven — applies rules to known facts to derive all consequences",
     "C) Heuristic-guided search through the hypothesis space",
     "D) Applicable only when the KB contains no rules"],
    "B",
    "FC = bottom-up / data-driven. Fires rules whose premises are satisfied and adds conclusions until the goal is derived or no rules fire.")

question(28,
    "Backward chaining is most efficient when:",
    ["A) You need to derive all possible consequences from the KB",
     "B) You need to prove one specific goal query",
     "C) The KB is very large and contains thousands of rules",
     "D) All rules have exactly one premise"],
    "B",
    "BC = goal-driven / top-down. Only explores rules relevant to the specific query — much more focused than FC.")

question(29,
    "Rule: 'If A AND B THEN C'. A is known to be true; B is unknown. Forward chaining will:",
    ["A) Fire the rule and conclude C, since A is true",
     "B) Not fire the rule — all premises must be simultaneously satisfied",
     "C) Add B to the KB as a new assumption to enable firing",
     "D) Ask the user to provide B interactively"],
    "B",
    "FC uses modus ponens: ALL premises must be true before a rule fires. A alone is insufficient for 'A AND B → C'.")

question(30,
    "MCV (Most Constrained Variable) in CSP selects:",
    ["A) The variable appearing in the most constraint expressions",
     "B) The variable with the FEWEST remaining legal values in its domain",
     "C) The value that eliminates the fewest choices for neighbouring variables",
     "D) The variable whose current assignment satisfies the most constraints"],
    "B",
    "MCV = Minimum Remaining Values (MRV) = fail-first heuristic. Fewest remaining values → most constrained → detect failure early.")

question(31,
    "LCV (Least Constraining Value) selects:",
    ["A) The value that eliminates the most choices for neighbouring variables",
     "B) The value that eliminates the FEWEST choices for neighbouring variables",
     "C) The variable with the fewest constraints in the graph",
     "D) The value that satisfies the largest number of hard constraints"],
    "B",
    "LCV = pick the value that leaves maximum flexibility for remaining variables. Tries to keep the search space as open as possible.")

question(32,
    "Arc consistency (AC-3) ensures that:",
    ["A) Every variable has exactly one remaining value",
     "B) For every value of variable X, there exists at least one compatible value in every constrained neighbour Y",
     "C) The constraint graph contains no cycles",
     "D) All constraints are binary (involving exactly 2 variables)"],
    "B",
    "Arc consistency: for arc (X,Y), remove values from X's domain that have no compatible value in Y. Prunes domains before search.")

question(33,
    "Which contrast between FC and BC is CORRECT?",
    ["A) FC is complete; BC is incomplete in all cases",
     "B) BC is better for proving a single specific goal; FC is better for deriving all consequences",
     "C) FC is goal-driven; BC is data-driven",
     "D) BC adds facts to the KB; FC only queries"],
    "B",
    "BC = goal-directed (single query). FC = derive everything (breadth). They complement each other.")

question(34,
    "The 'degree heuristic' for variable ordering in CSP breaks ties in MCV by selecting:",
    ["A) The variable with the smallest domain size",
     "B) The variable involved in the most constraints with UNASSIGNED variables",
     "C) The variable that was assigned earliest in the current path",
     "D) The variable with the highest-valued domain element"],
    "B",
    "Degree heuristic = most constrained by remaining neighbours. Used as tiebreaker when multiple variables have equal domain sizes.")

question(35,
    "In CSP, 'backtracking' means:",
    ["A) Running the search in reverse from the goal to the start",
     "B) Undoing the most recent variable assignment when a constraint is violated",
     "C) Restarting the entire search with a different variable ordering",
     "D) Applying arc consistency before assigning any variable"],
    "B",
    "Backtracking: assign one variable, check constraints; if violated, undo the assignment and try the next value.")

# ══════════════════════════════════════════════════════════
# L04: Local Search  (Q36-Q45)
# ══════════════════════════════════════════════════════════
question(36,
    "Hill climbing gets stuck because:",
    ["A) It uses too much memory tracking the full search history",
     "B) It only moves to strictly better neighbours and halts at local optima",
     "C) It generates too many successor states at each step",
     "D) It requires an admissible heuristic to function correctly"],
    "B",
    "HC accepts only improvements. At a local optimum (no better neighbour exists), it stops — even if the global optimum is far away.")

question(37,
    "Simulated Annealing accepts a WORSE solution (higher cost) with probability:",
    ["A) 1.0 always — it always accepts the worse move",
     "B) 0.0 always — it never accepts worse moves",
     "C) e^(−|Δcost|/T)",
     "D) Δcost / T"],
    "C",
    "SA acceptance: P = e^(−|Δ|/T). As T→large, P→1 (random walk). As T→0, P→0 (hill climbing).")

question(38,
    "As temperature T approaches 0 in Simulated Annealing, the algorithm converges to:",
    ["A) A random walk through the state space",
     "B) Breadth-first search",
     "C) Pure hill climbing — never accepts worse moves",
     "D) Genetic algorithm behaviour"],
    "C",
    "T→0 ⟹ e^(−|Δ|/T) → 0. Bad moves are rejected with probability approaching 1 → pure greedy hill climbing.")

question(39,
    "In a Genetic Algorithm, 'crossover' (recombination) does what?",
    ["A) Randomly flips one bit in a chromosome",
     "B) Combines segments of two parent chromosomes to create offspring",
     "C) Selects the top-K fittest individuals for the next generation",
     "D) Evaluates the fitness function on the current population"],
    "B",
    "Crossover: choose a cut point, swap tails of two parent chromosomes. Combines beneficial traits from both parents.")

question(40,
    "Which local search method is most suitable for finding the GLOBAL optimum on a complex multi-modal landscape?",
    ["A) Steepest-ascent hill climbing",
     "B) First-choice hill climbing",
     "C) Simulated Annealing with a very slow cooling schedule",
     "D) Single-restart hill climbing"],
    "C",
    "Slow SA cooling → more time at high T → thorough exploration → better chance of escaping local optima toward global.")

question(41,
    "In a Genetic Algorithm, the purpose of MUTATION is:",
    ["A) Combining beneficial traits from two parents",
     "B) Selecting high-fitness individuals for reproduction",
     "C) Introducing genetic diversity to prevent premature convergence to a local optimum",
     "D) Computing the fitness score of each individual"],
    "C",
    "Mutation randomly alters chromosomes, exploring new regions of the search space not reachable by crossover alone.")

question(42,
    "Hill climbing is NOT complete because:",
    ["A) It cannot handle graphs with cycles",
     "B) It can terminate at local optima without finding any solution (or the global optimum)",
     "C) It requires more memory than BFS",
     "D) It cannot handle continuous state spaces"],
    "B",
    "HC terminates when no better neighbour exists. This may be a local optimum, not the global — so it is neither complete nor optimal.")

question(43,
    "SA gets its name from which physical process?",
    ["A) Statistical averaging of molecular velocities",
     "B) Heating and slowly cooling metal to settle atoms into low-energy configurations",
     "C) The way neurons reduce firing rate after stimulation",
     "D) A cryptographic key scheduling algorithm"],
    "B",
    "In metallurgy, annealing = heat then slowly cool → crystal lattice settles into a globally minimal energy (optimal) state.")

question(44,
    "Which scenario is BEST suited for local search over systematic search?",
    ["A) Finding the shortest path in a small road network",
     "B) Proving a mathematical theorem from axioms",
     "C) Optimising a large scheduling problem where 'good enough' suffices",
     "D) Solving an 8-puzzle where the exact optimal solution is required"],
    "C",
    "Local search excels when: state space is huge, optimality is not strictly required, and memory is limited.")

question(45,
    "In SA minimisation, the acceptance probability for a move that INCREASES cost by Δ is e^(−Δ/T). If T=10 and Δ=5:",
    ["A) P ≈ 0.61",
     "B) P ≈ 0.37",
     "C) P ≈ 0.50",
     "D) P = 1.0 because Δ < T"],
    "A",
    "P = e^(−5/10) = e^(−0.5) ≈ 0.6065 ≈ 0.61. The move is accepted with 61% probability at this temperature.")

# ══════════════════════════════════════════════════════════
# L05: CSP & Games  (Q46-Q57)
# ══════════════════════════════════════════════════════════
question(46,
    "In a ZERO-SUM two-player game, if MAX gains utility of +5 in a state, MIN's utility in that state is:",
    ["A) +5",
     "B) 0",
     "C) −5",
     "D) Indeterminate — depends on the evaluation function"],
    "C",
    "Zero-sum: total utility is constant. If MAX gets +5, MIN gets −5. Their utilities always sum to zero.")

question(47,
    "In Minimax, a MIN node at depth 3 will:",
    ["A) Return the maximum of its children's values",
     "B) Return the minimum of its children's values",
     "C) Return the average of its children's values",
     "D) Return its heuristic h(n) value"],
    "B",
    "MIN player chooses the move that minimises MAX's score. MIN nodes return the minimum child value.")

question(48,
    "Alpha-Beta pruning cuts a branch at a MAX node when:",
    ["A) The node's value equals alpha",
     "B) The node's value ≥ beta (the best value the MIN ancestor can guarantee)",
     "C) The node has no children",
     "D) The heuristic is inadmissible"],
    "B",
    "At MAX node: if current value ≥ β, MIN will never choose this branch (has a better option). Prune remaining children.")

question(49,
    "With perfect move ordering, Alpha-Beta pruning reduces Minimax's O(b^d) to approximately:",
    ["A) O(b^{d/4})",
     "B) O(b^{2d/3})",
     "C) O(b^{d/2})",
     "D) O(d log b)"],
    "C",
    "Best case: O(b^{d/2}). Effective branching factor reduces to sqrt(b). Search depth effectively doubles for the same cost.")

question(50,
    "The 'alpha' variable in Alpha-Beta represents:",
    ["A) The best score found so far for MIN",
     "B) The best (highest) score found so far for MAX along the current path",
     "C) The evaluation function threshold",
     "D) The maximum search depth"],
    "B",
    "α = best value MAX can guarantee along current path (lower bound for MAX). β = best value MIN can guarantee (upper bound for MAX).")

question(51,
    "In CSP, a 'binary constraint' involves:",
    ["A) One variable with a binary (True/False) domain",
     "B) Exactly two variables",
     "C) A constraint with exactly two legal values",
     "D) Two unary constraints merged together"],
    "B",
    "Binary constraint: involves exactly 2 variables (e.g. X ≠ Y, X < Y). N-Queens has binary diagonal constraints.")

question(52,
    "Minimax assumes the opponent plays:",
    ["A) Randomly, with equal probability for each move",
     "B) Optimally — always choosing the move that minimises MAX's score",
     "C) Greedily — always choosing the locally best move",
     "D) As a cooperative partner trying to maximise joint utility"],
    "B",
    "Minimax worst-case assumption: opponent is a perfect rational player who always minimises your score.")

question(53,
    "In a game tree, an evaluation function is used because:",
    ["A) The full tree is always too deep to reach terminal states in time",
     "B) Alpha-Beta pruning cannot work without one",
     "C) Terminal states are always unreachable in practice",
     "D) It replaces the Minimax algorithm for shallow trees"],
    "A",
    "For complex games (chess, Go), computing to terminal states takes too long. Evaluation functions score non-terminal states heuristically.")

question(54,
    "Which CSP variable-ordering heuristic is used as a TIE-BREAKER for MCV?",
    ["A) LCV (Least Constraining Value)",
     "B) Arc consistency (AC-3)",
     "C) Degree heuristic (most constraints on unassigned variables)",
     "D) Forward checking"],
    "C",
    "When multiple variables are tied for fewest remaining values (MCV tie), use degree heuristic: pick the one constrained by the most unassigned neighbours.")

question(55,
    "Forward checking in CSP:",
    ["A) Checks arc consistency throughout the entire constraint graph",
     "B) After assigning a variable, removes inconsistent values only from directly constrained neighbours",
     "C) Applies backtracking before any variable is assigned",
     "D) Is equivalent to full AC-3 arc consistency"],
    "B",
    "Forward checking: after assigning X=v, remove values from neighbours that conflict with v. Simpler than full AC-3 but still prunes effectively.")

question(56,
    "Why does more pruning occur in Alpha-Beta with best-move ordering?",
    ["A) The tree becomes shallower",
     "B) Alpha and beta bounds tighten quickly when the best moves are explored first",
     "C) The evaluation function is computed fewer times",
     "D) MIN and MAX swap roles, eliminating redundant comparisons"],
    "B",
    "Best-first ordering: early exploration of strong moves sets tight α/β bounds → more subsequent branches get pruned.")

# ══════════════════════════════════════════════════════════
# L06: Machine Learning  (Q57-Q68)
# ══════════════════════════════════════════════════════════
question(57,
    "K-Means is an example of:",
    ["A) Supervised learning",
     "B) Reinforcement learning",
     "C) Unsupervised learning",
     "D) Semi-supervised learning"],
    "C",
    "K-Means has no labels — it discovers groupings purely from the data distribution. No teacher signal.")

question(58,
    "K-Means convergence is reached when:",
    ["A) The centroids reach the coordinate origin",
     "B) All cluster assignments stop changing between iterations",
     "C) The within-cluster variance drops below a threshold",
     "D) K equals the number of data points"],
    "B",
    "K-Means iterates: assign→update centroids→reassign. Stop when assignments are stable (centroids no longer move).")

question(59,
    "The entropy H(S) for a dataset with 4 positive and 4 negative examples (8 total) is:",
    ["A) 0.0",
     "B) 0.5",
     "C) 1.0",
     "D) 2.0"],
    "C",
    "H = −(4/8)log₂(4/8) − (4/8)log₂(4/8) = −0.5×(−1) − 0.5×(−1) = 1.0. 50/50 split = maximum entropy.")

question(60,
    "A Decision Tree feature with Information Gain = 0 means:",
    ["A) The feature perfectly separates the classes",
     "B) Splitting on that feature changes nothing — class distribution is the same in every sub-group",
     "C) The feature takes only one unique value",
     "D) The feature should be chosen as the root (zero error)"],
    "B",
    "IG=0 → weighted entropy after split equals entropy before split → the feature conveys no predictive information.")

question(61,
    "SVM finds the decision boundary that:",
    ["A) Minimises total training error",
     "B) Passes through the centroid of each class",
     "C) Maximises the margin (distance) between the nearest examples of each class",
     "D) Minimises the number of support vectors"],
    "C",
    "SVM = maximum-margin classifier. Margin = 2/||w||. Maximising margin = minimising ||w||.")

question(62,
    "Which statement about K-Means is FALSE?",
    ["A) K-Means always converges to a stable set of assignments",
     "B) K-Means is guaranteed to find the globally optimal clustering",
     "C) K-Means uses distance to the nearest centroid for assignment",
     "D) K-Means is sensitive to the initial centroid positions"],
    "B",
    "K-Means converges to a LOCAL optimum. Global optimality is not guaranteed. Different initialisations can give different results.")

question(63,
    "In Decision Tree learning, you create a LEAF node when:",
    ["A) The current node has more than 10 examples",
     "B) All examples at the node belong to the same class (entropy = 0)",
     "C) Information Gain exceeds 0.5 for the remaining features",
     "D) The tree depth exceeds the number of input features"],
    "B",
    "Pure node → entropy = 0 → no benefit in further splitting. Create a leaf labelled with that class.")

question(64,
    "Support vectors in SVM are:",
    ["A) All correctly classified training examples",
     "B) The training examples lying ON the margin boundary (closest to the hyperplane)",
     "C) Examples removed during preprocessing as outliers",
     "D) Examples that change the decision boundary if added or removed"],
    "B",
    "Support vectors: the critical examples that define the margin. Remove any non-support-vector and the hyperplane is unchanged.")

question(65,
    "The Decision Tree root should be:",
    ["A) The feature with the lowest Information Gain",
     "B) Always the first feature listed in the dataset",
     "C) The feature with the highest Information Gain on the full training set",
     "D) The feature with the most unique values"],
    "C",
    "Root = highest IG feature. Choosing the best splitter at each level greedily builds a compact, accurate tree.")

question(66,
    "K-Means with a poor initialisation will:",
    ["A) Never converge",
     "B) Converge to a local optimum that may be far from the best solution",
     "C) Produce clusters of equal size regardless of the data",
     "D) Automatically correct itself after enough iterations"],
    "B",
    "K-Means is sensitive to initial centroids. Poor start → bad local minimum. K-Means++ improves initialisation.")

question(67,
    "The SVM decision boundary w^T x + b = 0 is a:",
    ["A) Polynomial curve fitted to the training data",
     "B) Hyperplane separating the two classes",
     "C) Gaussian boundary modelling the class distributions",
     "D) Piecewise linear boundary defined by support vectors"],
    "B",
    "Linear SVM: decision boundary = hyperplane. In 2D, this is a line; in 3D, a plane; in d-D, a hyperplane.")

question(68,
    "A Decision Tree that perfectly fits the training data with H=0 at every leaf:",
    ["A) Is guaranteed to generalise well to new data",
     "B) Has zero training error but likely overfits — high variance on new data",
     "C) Is always the simplest possible tree for that dataset",
     "D) Cannot be constructed unless all features are binary"],
    "B",
    "All-pure leaves = zero training error = potential overfitting. Pruning or depth limits improve generalisation.")

# ══════════════════════════════════════════════════════════
# L07: Neural Networks  (Q69-Q80)
# ══════════════════════════════════════════════════════════
question(69,
    "Why can't the sign (step) activation function be used in backpropagation?",
    ["A) It produces outputs outside the range [0,1]",
     "B) Its derivative is zero almost everywhere, so gradients cannot flow back through it",
     "C) It is too slow to compute during forward pass",
     "D) It only supports binary inputs"],
    "B",
    "Backprop needs ∂output/∂net. Sign function derivative = 0 everywhere → gradient = 0 → weights never update → network cannot learn.")

question(70,
    "The sigmoid function σ(x) and its derivative are:",
    ["A) σ(x) = 1/(1+eˣ);   σ'(x) = σ(x)(1−σ(x))",
     "B) σ(x) = 1/(1+e^{−x});  σ'(x) = σ(x)(1−σ(x))",
     "C) σ(x) = 1/(1+e^{−x});  σ'(x) = 1 − σ(x)",
     "D) σ(x) = eˣ/(1+eˣ);  σ'(x) = σ(x)²"],
    "B",
    "σ(x) = 1/(1+e^{-x}). Its derivative is σ(x)·(1−σ(x)). Maximum value is 0.25 (at x=0). This maximum caps gradient magnitude.")

question(71,
    "The MLP output layer delta for output neuron p (with target y_p and output y_hat_p) is:",
    ["A) δ_p = y_p − y_hat_p",
     "B) δ_p = (y_p − y_hat_p) × y_hat_p × (1 − y_hat_p)",
     "C) δ_p = y_hat_p × (1 − y_hat_p)",
     "D) δ_p = Σ_j(W_jp × δ_j)"],
    "B",
    "Output delta = error signal × sigmoid derivative at that neuron. This is the chain-rule combination of loss gradient and activation gradient.")

question(72,
    "For a weight W_ij (input i → hidden j) in backprop, the update is:",
    ["A) W_new = W_old − η × δ_j × x_i",
     "B) W_new = W_old + η × δ_j × x_i",
     "C) W_new = W_old + η × (y − y_hat) × x_i",
     "D) W_new = W_old × (1 − η × δ_j)"],
    "B",
    "Weight update: W_new = W_old + η·δ_j·x_i. The delta already encodes error sign. Plus (+) because delta is defined to point toward the error reduction.")

question(73,
    "A perceptron CANNOT solve the XOR problem because:",
    ["A) The learning rate is too small",
     "B) XOR is not linearly separable — no single hyperplane can separate the classes",
     "C) The perceptron has too few weights",
     "D) XOR requires more than 2 input features"],
    "B",
    "XOR: (0,0)→0, (0,1)→1, (1,0)→1, (1,1)→0. No straight line separates 0s from 1s. Need MLP with hidden layer.")

question(74,
    "In backprop, the weight connecting input x₂=0 to any hidden neuron does NOT change because:",
    ["A) The hidden neuron's activation was zero",
     "B) ΔW = η·δ·x_i = η·δ·0 = 0 — the update is zero when the input is zero",
     "C) The output error was zero",
     "D) The learning rate η was set to zero"],
    "B",
    "Update ΔW ∝ x_i. When x_i=0, the weight's gradient is zero regardless of the error signal. Zero input → zero update.")

question(75,
    "Mini-batch gradient descent with N=10,000 examples and batch size B=200 runs how many weight updates per EPOCH?",
    ["A) 10,000",
     "B) 200",
     "C) 50",
     "D) 2,000,000"],
    "C",
    "Updates per epoch = N/B = 10,000/200 = 50. Each mini-batch produces one gradient estimate and one weight update.")

question(76,
    "The vanishing gradient problem occurs in deep networks because:",
    ["A) Weights grow too large, saturating neurons",
     "B) Sigmoid derivatives are ≤ 0.25, so backpropagating through many layers shrinks gradients toward zero",
     "C) The learning rate is too high for deep networks",
     "D) The training dataset is too small for the model capacity"],
    "B",
    "Each sigmoid layer multiplies gradient by σ'(x) ≤ 0.25. With L layers: gradient ≤ (0.25)^L → effectively zero for large L. ReLU and LSTMs address this.")

question(77,
    "Stochastic Gradient Descent (SGD) vs Batch Gradient Descent: which is TRUE?",
    ["A) Batch GD uses one example per update; SGD uses the full dataset",
     "B) SGD uses one example per update; noisier but faster per parameter update",
     "C) Both compute identical gradients — the names are interchangeable",
     "D) SGD only works when the loss surface is convex"],
    "B",
    "SGD: one example → one update. Noisy gradient estimate but very fast. Batch: full dataset → exact gradient but slow per update.")

question(78,
    "The bias weight in a neuron serves the purpose of:",
    ["A) Normalising the output to lie in [0,1]",
     "B) Shifting the activation threshold independently of the inputs",
     "C) Regularising the weights to prevent overfitting",
     "D) Scaling the learning rate during backpropagation"],
    "B",
    "Bias = intercept. Without bias, decision boundaries must pass through the origin, severely limiting the model's flexibility.")

question(79,
    "Perceptron learning: a weight update occurs ONLY when:",
    ["A) The perceptron correctly classifies the example",
     "B) The perceptron misclassifies the example (y ≠ ŷ)",
     "C) The dot product w·x exceeds the threshold",
     "D) The learning rate η > 0.5"],
    "B",
    "Perceptron rule: Δw = η·(y − ŷ)·x. When y = ŷ, the error is 0 → Δw = 0. Only update on mistakes.")

question(80,
    "Which of the following is NOT a standard variant of gradient descent?",
    ["A) Batch gradient descent",
     "B) Stochastic gradient descent",
     "C) Mini-batch gradient descent",
     "D) Diagonal gradient descent"],
    "D",
    "Batch, SGD, and mini-batch are the three canonical variants. 'Diagonal gradient descent' is not a recognised optimisation method.")

# ══════════════════════════════════════════════════════════
# L08: Deep Learning  (Q81-Q90)
# ══════════════════════════════════════════════════════════
question(81,
    "A convolution with a 6×6 input image, 3×3 kernel, stride=1, no padding produces an output of size:",
    ["A) 6×6",
     "B) 4×4",
     "C) 3×3",
     "D) 2×2"],
    "B",
    "Output = (6−3)/1 + 1 = 4. So 4×4.")

question(82,
    "Max pooling with a 2×2 window and stride 2 applied to a 4×4 feature map produces:",
    ["A) A 4×4 output — dimensions are unchanged",
     "B) A 3×3 output",
     "C) A 2×2 output — dimensions are halved",
     "D) A 1×1 output — all spatial information is discarded"],
    "C",
    "Output = (4−2)/2 + 1 = 2. So 2×2. Max pool 2×2 stride 2 halves each spatial dimension.")

question(83,
    "TRANSLATION INVARIANCE in CNNs is primarily provided by:",
    ["A) Convolutional layers (via weight sharing)",
     "B) Fully connected layers",
     "C) Pooling layers",
     "D) Batch normalisation"],
    "C",
    "Pooling discards precise spatial positions → the same feature anywhere in the image produces the same pooled response.")

question(84,
    "An RNN suffers from the 'vanishing gradient' problem because:",
    ["A) It uses fully connected layers which have too many parameters",
     "B) Gradients are multiplied by the recurrent weight matrix at each time step, shrinking exponentially over long sequences",
     "C) It applies max pooling which destroys gradient information",
     "D) Its hidden state size is fixed, causing information bottleneck"],
    "B",
    "At each step, gradients are multiplied through the recurrent weights. Many steps → gradients vanish (or explode) → RNN cannot model long-range dependencies.")

question(85,
    "In an LSTM, the FORGET gate output close to 0 means:",
    ["A) The network retains the entire previous cell state",
     "B) The network erases (forgets) that portion of the cell state",
     "C) The network outputs the full cell state value",
     "D) The new input is ignored at this time step"],
    "B",
    "Forget gate: C_t = f_t ⊙ C_{t-1} + i_t ⊙ g_t. f_t≈0 → C_{t-1} term → 0 → cell state is erased for that unit.")

question(86,
    "In a GAN, training reaches equilibrium when:",
    ["A) The generator loss drops to zero",
     "B) The discriminator outputs 0.5 for both real and generated samples",
     "C) The discriminator correctly classifies 100% of samples",
     "D) The generator exactly memorises the training images"],
    "B",
    "Equilibrium: G produces indistinguishable samples → D cannot tell real from fake → outputs 0.5 (random guess).")

question(87,
    "Which CNN layer has LEARNABLE PARAMETERS (weights)?",
    ["A) Max pooling layer",
     "B) ReLU activation layer",
     "C) Convolutional layer (filters)",
     "D) Dropout layer"],
    "C",
    "Convolutional filters are the learnable parameters. Pooling, ReLU, and Dropout are parameter-free operations.")

question(88,
    "The fully connected (FC) layer at the end of a CNN does what?",
    ["A) Applies a spatial convolution filter to the final feature maps",
     "B) Flattens all feature maps into a vector and learns a global classification mapping",
     "C) Performs max pooling across all channels simultaneously",
     "D) Reduces the number of filters by a factor of 2"],
    "B",
    "FC layer: flatten spatial feature maps → dense connections → class score logits. Combines global spatial information for the final prediction.")

question(89,
    "For crowd counting using deep learning, CNN is preferred over a plain MLP because:",
    ["A) CNNs use fewer parameters in total",
     "B) CNNs exploit spatial locality — convolutional filters detect heads/bodies regardless of their position via weight sharing",
     "C) MLPs cannot process images larger than 28×28 pixels",
     "D) CNNs are always faster to train than MLPs on image data"],
    "B",
    "CNNs: weight sharing + local receptive field + translation invariance via pooling. Critical for images where the same pattern appears at different positions.")

question(90,
    "The three LSTM gates control:",
    ["A) Input size, hidden size, and output size",
     "B) What to forget from cell state, what new info to write, and what to output",
     "C) Learning rate, momentum, and weight decay",
     "D) Encoder, decoder, and attention weights"],
    "B",
    "LSTM gates: Forget (erase cell state), Input (write new info), Output (expose cell state as hidden state). Together they enable selective long-term memory.")

# ══════════════════════════════════════════════════════════
# L09: Bayesian Networks & Naive Bayes  (Q91-Q100)
# ══════════════════════════════════════════════════════════
question(91,
    "In Bayes' theorem P(H|E) = P(E|H)·P(H)/P(E), the term P(H) is called:",
    ["A) The likelihood",
     "B) The posterior",
     "C) The prior",
     "D) The marginal likelihood"],
    "C",
    "Prior P(H) = your belief BEFORE seeing evidence. Posterior P(H|E) = updated belief AFTER evidence. Likelihood P(E|H) = how probable is E given H.")

question(92,
    "In a Bayesian Network, an arrow from node A to node B signifies:",
    ["A) A and B are conditionally independent given all other nodes",
     "B) A directly influences B — B's probability depends on the value of A",
     "C) B is more probable than A",
     "D) A and B must always have the same value"],
    "B",
    "Arrows in a BN = direct probabilistic dependence. A→B means B's CPT is conditioned on A (A is a cause/parent of B).")

question(93,
    "Naive Bayes makes the 'naive' assumption of:",
    ["A) Equal prior probability for all classes",
     "B) Conditional independence of features given the class label",
     "C) Normally distributed features within each class",
     "D) All features having exactly two possible values"],
    "B",
    "Naive = features are conditionally independent given class: P(f1,...,fn|C) = ∏P(fi|C). This allows efficient computation but may be wrong in practice.")

question(94,
    "In the Naive Bayes DAG, the class node is:",
    ["A) A child of all feature nodes (features cause the class)",
     "B) The parent of all feature nodes (class causes the features)",
     "C) Isolated — not connected to any feature node",
     "D) Connected to only the most predictive feature"],
    "B",
    "NB DAG: Class is the root. Arrows point FROM class TO each feature. This encodes: the class generates (causes) the observed feature values.")

question(95,
    "In Naive Bayes classification, you compare the scores for each class and predict:",
    ["A) The class whose score exceeds 0.5",
     "B) The class with the highest score (argmax) — no normalisation needed",
     "C) The class whose score is closest to the prior probability",
     "D) The average of all class scores"],
    "B",
    "Argmax rule: predict the class with highest Score(C) = P(C)·∏P(fi|C). Scores need not be normalised for prediction.")

question(96,
    "To compute P(C) in a Bayesian Network where C has parents A and B, you:",
    ["A) Use only the CPT entry P(C|A=True, B=True)",
     "B) Sum over all combinations of parent values: Σ_a Σ_b P(C|A=a,B=b)·P(A=a)·P(B=b)",
     "C) Multiply P(C|A) by P(C|B) and divide by P(C)",
     "D) Read P(C) directly from the CPT without marginalisation"],
    "B",
    "Marginalise over all parent configurations: P(C) = Σ over all parent combos of [P(C|parents) × P(parents)].")

question(97,
    "A COMMON MISTAKE in Naive Bayes: with 4 Low Risk training instances (IDs 2,4,5,6), where 2 have High income (IDs 4,6), P(Income=High | Low Risk) = ?",
    ["A) 3/4 — counting incorrectly",
     "B) 2/4 = 0.5 — correct count",
     "C) 1/4 — inverse count",
     "D) 4/7 — using total instances"],
    "B",
    "P(Income=High|LR) = count of LR instances with High income / total LR instances = 2/4 = 0.5. The classic error is using 3/4 (accidentally including a HR instance).")

question(98,
    "P(W|C) = P(C|W)·P(W)/P(C). If P(C|W=True)=0.762, P(W=True)=0.6, P(C)=0.5588, then P(W=True|C) ≈:",
    ["A) 0.762",
     "B) 0.600",
     "C) 0.818",
     "D) 0.457"],
    "C",
    "P(W|C) = 0.762 × 0.6 / 0.5588 = 0.4572 / 0.5588 ≈ 0.818. Bayes' theorem: multiply likelihood × prior, divide by evidence.")

question(99,
    "In a Bayesian Network, 'diagnostic inference' (effect → cause) is:",
    ["A) P(cause | effect) — computed using Bayes' theorem over the network",
     "B) P(effect | cause) — read directly from the CPT",
     "C) Only possible when the network has no hidden nodes",
     "D) The same computation as predictive inference"],
    "A",
    "Diagnostic (bottom-up): given observed effects, infer P(cause|evidence) using Bayes' rule. Predictive (top-down): P(effect|cause) is read from CPT directly.")

question(100,
    "The Naive Bayes Score(HR) = 3/7 × 2/3 × 2/3 × 1/3 = 0.0635 and Score(LR) = 4/7 × 2/4 × 1/4 × 3/4 = 0.0536. The final classification is:",
    ["A) Low Risk — Score(LR) is derived from more training instances",
     "B) High Risk — highest score wins with argmax rule",
     "C) Undecided — the scores must be normalised to probabilities first",
     "D) High Risk only if Score(HR) > 0.5 after normalisation"],
    "B",
    "Argmax: compare raw scores. Score(HR)=0.0635 > Score(LR)=0.0536 → predict HIGH RISK. No normalisation needed for classification.")

# ══════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════
doc.add_heading('COMP4431 — 100 MC Warm-Up Quiz', 0).runs[0].font.color.rgb = RGBColor(0x1D, 0x40, 0xAF)

p = doc.add_paragraph('Covers L01–L09. Time target: 60–75 minutes (≈45 sec/question). Answer sheet + explanations on page 2.')
p.runs[0].font.italic = True
p.runs[0].font.size = Pt(10)

doc.add_paragraph().paragraph_format.space_after = Pt(0)

SECTIONS = [
    (1,  10,  'L01: Introduction to AI'),
    (11, 25,  'L02: Heuristic Search'),
    (26, 35,  'L03: Knowledge-Based Agents'),
    (36, 45,  'L04: Local Search'),
    (46, 56,  'L05: CSP & Games'),
    (57, 68,  'L06: Machine Learning'),
    (69, 80,  'L07: Neural Networks'),
    (81, 90,  'L08: Deep Learning'),
    (91, 100, 'L09: Bayesian Networks & Naive Bayes'),
]

# ── Print questions ────────────────────────────────────────
for (sec_start, sec_end, sec_title) in SECTIONS:
    h2(sec_title)
    for (num, stem, options, correct, explanation) in QUESTIONS:
        if sec_start <= num <= sec_end:
            # Question stem
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(f'Q{num}. {stem}')
            run.font.bold = True
            run.font.size = Pt(10.5)
            # Options
            for opt in options:
                p2 = doc.add_paragraph(opt, style='List Bullet')
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after  = Pt(0)
                p2.paragraph_format.left_indent  = Cm(0.5)
                for r in p2.runs:
                    r.font.size = Pt(10)

doc.add_page_break()

# ── Answer key + explanations ─────────────────────────────
h1('Answer Key & Explanations')

p = doc.add_paragraph()
run = p.add_run('Quick answer grid:')
run.font.bold = True
run.font.size = Pt(10)

# 10-column grid table
grid_data = []
for i in range(0, 100, 10):
    row = []
    for j in range(10):
        idx = i + j
        if idx < len(QUESTIONS):
            num, _, _, correct, _ = QUESTIONS[idx]
            row.append(f'Q{num}: {correct}')
    grid_data.append(row)

tbl = doc.add_table(rows=10, cols=10)
tbl.style = 'Table Grid'
for r_idx, row in enumerate(grid_data):
    for c_idx, cell_text in enumerate(row):
        cell = tbl.cell(r_idx, c_idx)
        cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8.5)
                run.font.bold = True
        # colour correct-answer letter in green
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Full explanations
h2('Detailed Explanations')
for (num, stem, options, correct, explanation) in QUESTIONS:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f'Q{num} [{correct}]  ')
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)  # green
    r2 = p.add_run(explanation)
    r2.font.size = Pt(9.5)

doc.save('COMP4431_MC_Quiz_100Q.docx')
print('Saved COMP4431_MC_Quiz_100Q.docx')
